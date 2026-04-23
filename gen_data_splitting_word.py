from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
title = doc.add_heading('MODULE 2.2: PHÂN TÁCH DỮ LIỆU (DATA SPLITTING)', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('1. Giới Thiệu', 1)
doc.add_paragraph('Module Data Splitting mô tả cách chia dữ liệu thành train & test sets.')
doc.add_paragraph('Mục đích: Đánh giá chính xác performance của model trên dữ liệu chưa biết.')
doc.add_paragraph('Tỷ lệ chuẩn: 80% training, 20% testing', style='List Bullet')

doc.add_heading('2. 3.3.1: Tập Huấn Luyện (Train Set - 80%)', 2)
doc.add_paragraph('Định nghĩa:', style='List Bullet')
doc.add_paragraph('80% dữ liệu được dùng để huấn luyện model', style='List Bullet 2')
doc.add_paragraph('Model học patterns từ dữ liệu này', style='List Bullet 2')
doc.add_paragraph('Kích thước:', style='List Bullet')
doc.add_paragraph('Tổng dữ liệu: 1200 mẫu', style='List Bullet 2')
doc.add_paragraph('Train set: 1200 × 0.80 = 960 mẫu', style='List Bullet 2')
doc.add_paragraph('Các mẫu:', style='List Bullet')
doc.add_paragraph('~80 mẫu per major × 12 majors', style='List Bullet 2')
doc.add_paragraph('Balanced distribution', style='List Bullet 2')

doc.add_heading('2.1 Cách Tạo Train Set', 2)
doc.add_paragraph('Bước 1: Load cleaned data', style='List Number')
doc.add_paragraph('df = pd.read_csv("data/raw/students.csv")', style='List Bullet 2')
doc.add_paragraph('Bước 2: Shuffle data (random order)', style='List Number')
doc.add_paragraph('df = df.sample(frac=1, random_state=42)', style='List Bullet 2')
doc.add_paragraph('Bước 3: Split 80%', style='List Number')
doc.add_paragraph('train_idx = int(0.8 * len(df))', style='List Bullet 2')
doc.add_paragraph('X_train = df[:train_idx]', style='List Bullet 2')
doc.add_paragraph('Bước 4: Extract features & labels', style='List Number')
doc.add_paragraph('X_train = train_df.drop("major", axis=1)', style='List Bullet 2')
doc.add_paragraph('y_train = train_df["major"]', style='List Bullet 2')

doc.add_heading('2.2 Ưu Điểm Train Set 80%', 2)
doc.add_paragraph('✓ Đủ dữ liệu để model học patterns', style='List Bullet')
doc.add_paragraph('✓ Reduce overfitting risk', style='List Bullet')
doc.add_paragraph('✓ Balanced representation của mỗi class', style='List Bullet')
doc.add_paragraph('✓ Model có chance tốt để converge', style='List Bullet')

doc.add_heading('2.3 Thống Kê Train Set', 2)
table1 = doc.add_table(rows=4, cols=3)
table1.style = 'Light Grid Accent 1'
table1.rows[0].cells[0].text = 'Metric'
table1.rows[0].cells[1].text = 'Value'
table1.rows[0].cells[2].text = 'Note'
table1.rows[1].cells[0].text = 'Total Samples'
table1.rows[1].cells[1].text = '960'
table1.rows[1].cells[2].text = '80% of 1200'
table1.rows[2].cells[0].text = 'Per Major'
table1.rows[2].cells[1].text = '~64'
table1.rows[2].cells[2].text = '960 ÷ 15 majors'
table1.rows[3].cells[0].text = 'Features'
table1.rows[3].cells[1].text = '~120'
table1.rows[3].cells[2].text = 'OneHot + TF-IDF'

doc.add_heading('3. 3.3.2: Tập Kiểm Thử (Test Set - 20%)', 2)
doc.add_paragraph('Định nghĩa:', style='List Bullet')
doc.add_paragraph('20% dữ liệu được dùng để đánh giá model', style='List Bullet 2')
doc.add_paragraph('Model chưa bao giờ thấy dữ liệu này', style='List Bullet 2')
doc.add_paragraph('Kích thước:', style='List Bullet')
doc.add_paragraph('Tổng dữ liệu: 1200 mẫu', style='List Bullet 2')
doc.add_paragraph('Test set: 1200 × 0.20 = 240 mẫu', style='List Bullet 2')
doc.add_paragraph('Các mẫu:', style='List Bullet')
doc.add_paragraph('~16 mẫu per major × 15 majors', style='List Bullet 2')
doc.add_paragraph('Balanced distribution', style='List Bullet 2')

doc.add_heading('3.1 Cách Tạo Test Set', 2)
doc.add_paragraph('Bước 1: Load cleaned data (same as train)', style='List Number')
doc.add_paragraph('df = pd.read_csv("data/raw/students.csv")', style='List Bullet 2')
doc.add_paragraph('Bước 2: Shuffle data (random_state=42)', style='List Number')
doc.add_paragraph('df = df.sample(frac=1, random_state=42)', style='List Bullet 2')
doc.add_paragraph('Bước 3: Split 20%', style='List Number')
doc.add_paragraph('test_idx = int(0.8 * len(df))', style='List Bullet 2')
doc.add_paragraph('X_test = df[test_idx:]', style='List Bullet 2')
doc.add_paragraph('Bước 4: Extract features & labels', style='List Number')
doc.add_paragraph('X_test = test_df.drop("major", axis=1)', style='List Bullet 2')
doc.add_paragraph('y_test = test_df["major"]', style='List Bullet 2')

doc.add_heading('3.2 Ưu Điểm Test Set 20%', 2)
doc.add_paragraph('✓ Độc lập từ training data', style='List Bullet')
doc.add_paragraph('✓ Đánh giá generalization ability', style='List Bullet')
doc.add_paragraph('✓ Detect overfitting', style='List Bullet')
doc.add_paragraph('✓ Cân đủ lớn (240 samples)', style='List Bullet')

doc.add_heading('3.3 Thống Kê Test Set', 2)
table2 = doc.add_table(rows=4, cols=3)
table2.style = 'Light Grid Accent 1'
table2.rows[0].cells[0].text = 'Metric'
table2.rows[0].cells[1].text = 'Value'
table2.rows[0].cells[2].text = 'Note'
table2.rows[1].cells[0].text = 'Total Samples'
table2.rows[1].cells[1].text = '240'
table2.rows[1].cells[2].text = '20% of 1200'
table2.rows[2].cells[0].text = 'Per Major'
table2.rows[2].cells[1].text = '~16'
table2.rows[2].cells[2].text = '240 ÷ 15 majors'
table2.rows[3].cells[0].text = 'Features'
table2.rows[3].cells[1].text = '~120'
table2.rows[3].cells[2].text = 'OneHot + TF-IDF'

doc.add_heading('4. So Sánh Train vs Test', 1)
table3 = doc.add_table(rows=7, cols=3)
table3.style = 'Light Grid Accent 1'
table3.rows[0].cells[0].text = 'Đặc Điểm'
table3.rows[0].cells[1].text = 'Train Set (80%)'
table3.rows[0].cells[2].text = 'Test Set (20%)'
table3.rows[1].cells[0].text = 'Kích Thước'
table3.rows[1].cells[1].text = '960 samples'
table3.rows[1].cells[2].text = '240 samples'
table3.rows[2].cells[0].text = 'Mục Đích'
table3.rows[2].cells[1].text = 'Huấn luyện model'
table3.rows[2].cells[2].text = 'Đánh giá performance'
table3.rows[3].cells[0].text = 'Model Nhìn Thấy'
table3.rows[3].cells[1].text = 'Có'
table3.rows[3].cells[2].text = 'Không'
table3.rows[4].cells[0].text = 'Bias'
table3.rows[4].cells[1].text = 'Cao (memorization)'
table3.rows[4].cells[2].text = 'Thấp (generalization)'
table3.rows[5].cells[0].text = 'Per Major'
table3.rows[5].cells[1].text = '~64 samples'
table3.rows[5].cells[2].text = '~16 samples'
table3.rows[6].cells[0].text = 'Random State'
table3.rows[6].cells[1].text = '42 (reproducible)'
table3.rows[6].cells[2].text = '42 (same as train)'

doc.add_heading('5. Python Implementation', 1)
doc.add_paragraph('from sklearn.model_selection import train_test_split', style='List Bullet')
doc.add_paragraph('X_train, X_test, y_train, y_test = train_test_split(', style='List Bullet')
doc.add_paragraph('X, y,', style='List Bullet 2')
doc.add_paragraph('test_size=0.2,', style='List Bullet 2')
doc.add_paragraph('random_state=42,', style='List Bullet 2')
doc.add_paragraph('stratify=y  # Giữ balanced distribution', style='List Bullet 2')
doc.add_paragraph(')', style='List Bullet')

doc.add_heading('6. Stratification (Quan Trọng!)', 1)
doc.add_paragraph('Vấn đề:', style='List Bullet')
doc.add_paragraph('Random split có thể mất balance', style='List Bullet 2')
doc.add_paragraph('Ví dụ: Test set có quá nhiều một class', style='List Bullet 2')
doc.add_paragraph('Giải pháp:', style='List Bullet')
doc.add_paragraph('Sử dụng stratify=y trong train_test_split', style='List Bullet 2')
doc.add_paragraph('Đảm bảo train & test có same distribution', style='List Bullet 2')
doc.add_paragraph('Mỗi major chiếm ~6.7% (1/15) trong cả train & test', style='List Bullet 2')

doc.add_heading('7. Kết Luận', 1)
doc.add_paragraph('✓ 80/20 split là chuẩn trong ML', style='List Bullet')
doc.add_paragraph('✓ Train set để model học patterns', style='List Bullet')
doc.add_paragraph('✓ Test set để đánh giá generalization', style='List Bullet')
doc.add_paragraph('✓ Stratification giữ balance distribution', style='List Bullet')
doc.add_paragraph('✓ Random state=42 giúp reproducible', style='List Bullet')

doc.save('DATA_SPLITTING_MODULE.docx')
print("✅ Data Splitting Module completed: DATA_SPLITTING_MODULE.docx")
