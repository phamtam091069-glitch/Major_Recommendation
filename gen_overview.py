#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('HỆ THỐNG AI TƯ VẤN NGÀNH HỌC', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_paragraph('TỔNG QUAN TOÀN DIỆN')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 1. Basic Info
doc.add_heading('1. THÔNG TIN CƠ BẢN', level=1)
info_table = doc.add_table(rows=7, cols=2)
info_table.style = 'Light Grid Accent 1'
rows = info_table.rows
rows[0].cells[0].text = 'Tên hệ thống'
rows[0].cells[1].text = 'Hệ thống AI tư vấn ngành học phù hợp'
rows[1].cells[0].text = 'Công nghệ'
rows[1].cells[1].text = 'Flask (Python) + Machine Learning + NLP'
rows[2].cells[0].text = 'Mục đích'
rows[2].cells[1].text = 'Gợi ý Top 3 ngành đại học phù hợp'
rows[3].cells[0].text = 'Tổng ngành'
rows[3].cells[1].text = '74 ngành đại học'
rows[4].cells[0].text = 'Phạm vi'
rows[4].cells[1].text = 'Toàn bộ lĩnh vực giáo dục đại học'
rows[5].cells[0].text = 'Ngôn ngữ'
rows[5].cells[1].text = 'Tiếng Việt (với dấu + không dấu)'
rows[6].cells[0].text = 'Trạng thái'
rows[6].cells[1].text = 'Đã triển khai & Hoạt động'

# 2. Architecture
doc.add_heading('2. KIẾN TRÚC HỆ THỐNG (3-TIER)', level=1)

doc.add_heading('TIER 1: CLIENT (Frontend)', level=2)
doc.add_paragraph('HTML/CSS/JavaScript', style='List Bullet')
doc.add_paragraph('Form nhập liệu 8 trường (Profile học sinh)', style='List Bullet')
doc.add_paragraph('Chuẩn hóa tiếng Việt (Xóa dấu)', style='List Bullet')
doc.add_paragraph('Giao diện hiển thị kết quả Top 3', style='List Bullet')
doc.add_paragraph('Files: templates/index.html, static/script.js, static/style.css', style='List Bullet')

doc.add_heading('TIER 2: SERVER (Backend)', level=2)
doc.add_paragraph('Flask Web Server (Python)', style='List Bullet')
doc.add_paragraph('REST API: GET /, POST /predict, POST /chat, GET /health', style='List Bullet')
doc.add_paragraph('Xử lý dự đoán + Chatbot AI', style='List Bullet')
doc.add_paragraph('Files: app.py, utils/predictor.py, utils/chatbot.py', style='List Bullet')

doc.add_heading('TIER 3: STORAGE (Data Layer)', level=2)
doc.add_paragraph('Models: rf_model.pkl, ohe.pkl, tfidf.pkl, classes.pkl', style='List Bullet')
doc.add_paragraph('Config: hybrid_config.json, majors.json', style='List Bullet')
doc.add_paragraph('Data: students.csv, majors_profiles.json, salary_benchmarks.json', style='List Bullet')

# 3. Input Fields
doc.add_heading('3. ĐẦU VÀO NGƯỜI DÙNG (8 TRƯỜNG)', level=1)
input_table = doc.add_table(rows=9, cols=3)
input_table.style = 'Light Grid Accent 1'
rows = input_table.rows
rows[0].cells[0].text = 'Loại'
rows[0].cells[1].text = 'Trường'
rows[0].cells[2].text = 'Mô tả'
rows[1].cells[0].text = 'Bắt buộc'
rows[1].cells[1].text = 'Sở thích chính'
rows[1].cells[2].text = 'Lĩnh vực quan tâm chính'
rows[2].cells[0].text = 'Bắt buộc'
rows[2].cells[1].text = 'Môn học yêu thích'
rows[2].cells[2].text = 'Môn học yêu thích nhất'
rows[3].cells[0].text = 'Bắt buộc'
rows[3].cells[1].text = 'Tính cách'
rows[3].cells[2].text = 'Tính cách của học sinh'
rows[4].cells[0].text = 'Bắt buộc'
rows[4].cells[1].text = 'Kỹ năng nổi bật'
rows[4].cells[2].text = 'Kỹ năng mạnh nhất'
rows[5].cells[0].text = 'Bắt buộc'
rows[5].cells[1].text = 'Môi trường làm việc'
rows[5].cells[2].text = 'Môi trường mơ ước'
rows[6].cells[0].text = 'Bắt buộc'
rows[6].cells[1].text = 'Mục tiêu nghề nghiệp'
rows[6].cells[2].text = 'Mục tiêu dự kiến'
rows[7].cells[0].text = 'Tùy chọn'
rows[7].cells[1].text = 'Mô tả bản thân'
rows[7].cells[2].text = 'Thông tin bổ sung'
rows[8].cells[0].text = 'Tùy chọn'
rows[8].cells[1].text = 'Định hướng tương lai'
rows[8].cells[2].text = 'Kế hoạch tương lai dài hạn'

# 4. Scoring Formula
doc.add_heading('4. CÔNG THỨC TÍNH ĐIỂM (HYBRID)', level=1)
doc.add_paragraph('Final_Score = 0.30 × Model_Score + 0.70 × Criteria_Score + Rule_Boost')
doc.add_paragraph('Kết quả: 0-100 điểm')

doc.add_heading('Model Score (30% trọng số)', level=2)
doc.add_paragraph('RandomForest Classifier (Calibrated)', style='List Bullet')
doc.add_paragraph('TF-IDF Cosine Similarity', style='List Bullet')
doc.add_paragraph('Công thức: 60% ML Probability + 40% Cosine Similarity', style='List Bullet')

doc.add_heading('Criteria Score (70% trọng số) - 8 Tiêu chí', level=2)
criteria = [
    '1. Sở thích chính: 23%',
    '2. Định hướng tương lai: 20%',
    '3. Kỹ năng nổi bật: 16%',
    '4. Tính cách: 14%',
    '5. Môi trường làm việc: 12%',
    '6. Môn học yêu thích: 8%',
    '7. Mô tả bản thân: 4%',
    '8. Mục tiêu nghề nghiệp: 3%'
]
for c in criteria:
    doc.add_paragraph(c, style='List Bullet')

# 5. API Endpoints
doc.add_heading('5. CÁC API CHÍNH', level=1)
api_table = doc.add_table(rows=5, cols=4)
api_table.style = 'Light Grid Accent 1'
rows = api_table.rows
rows[0].cells[0].text = 'Endpoint'
rows[0].cells[1].text = 'Method'
rows[0].cells[2].text = 'Mục đích'
rows[0].cells[3].text = 'Input/Output'
rows[1].cells[0].text = '/'
rows[1].cells[1].text = 'GET'
rows[1].cells[2].text = 'Trang form'
rows[1].cells[3].text = 'HTML form'
rows[2].cells[0].text = '/predict'
rows[2].cells[1].text = 'POST'
rows[2].cells[2].text = 'Dự đoán ngành'
rows[2].cells[3].text = 'JSON → Top 3 + Scores'
rows[3].cells[0].text = '/health'
rows[3].cells[1].text = 'GET'
rows[3].cells[2].text = 'Kiểm tra trạng thái'
rows[3].cells[3].text = 'Status + Ready'
rows[4].cells[0].text = '/chat'
rows[4].cells[1].text = 'POST'
rows[4].cells[2].text = 'Chatbot AI'
rows[4].cells[3].text = 'Question → Response'

# 6. Output
doc.add_heading('6. KẾT QUẢ ĐẦU RA', level=1)
doc.add_paragraph('Mỗi dự đoán trả về Top 3 ngành với:')
doc.add_paragraph('Tên ngành (Tiếng Việt & không dấu)', style='List Bullet')
doc.add_paragraph('Điểm cuối (0-100)', style='List Bullet')
doc.add_paragraph('Điểm từ ML model', style='List Bullet')
doc.add_paragraph('Điểm từ tiêu chí minh bạch', style='List Bullet')
doc.add_paragraph('Độ tin cậy (Cao/Trung bình/Tham khảo)', style='List Bullet')
doc.add_paragraph('Lời khuyên & gợi ý chi tiết', style='List Bullet')

# 7. Features
doc.add_heading('7. ĐIỂM NỔI VẬY CỦA HỆ THỐNG', level=1)
features = [
    'Hybrid Approach: Kết hợp ML + Content-based + Rule-based',
    'Minh bạch: Giải thích từng bước tính toán (8 tiêu chí)',
    'Linh hoạt: Hỗ trợ 74 ngành đa lĩnh vực',
    'Đa ngôn ngữ: Tiếng Việt đầy đủ (với dấu + không dấu)',
    'AI Chatbot: Tương tác thêm để hỗ trợ quyết định',
    'Độ tin cậy: Tính confidence score & mức độ phù hợp',
    'Mở rộng: Dễ thêm ngành mới hoặc điều chỉnh trọng số',
    'Tested: Có test suites cho quality assurance'
]
for f in features:
    doc.add_paragraph('✓ ' + f, style='List Bullet')

# Save
output_path = 'c:/Users/huyen/Downloads/major-recommendation/SYSTEM_OVERVIEW.docx'
doc.save(output_path)
print(f'✅ System Overview Word document created: {output_path}')
