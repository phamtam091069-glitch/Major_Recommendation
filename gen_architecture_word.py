from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('MODULE 4: KIẾN TRÚC HỆ THỐNG (SYSTEM ARCHITECTURE)', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Introduction
doc.add_heading('1. Giới Thiệu', 1)
doc.add_paragraph('Hệ thống tư vấn ngành học được xây dựng theo kiến trúc 3 tầng (3-Tier Architecture):')
doc.add_paragraph('Frontend Layer: Giao diện người dùng (HTML, CSS, JavaScript)', style='List Bullet')
doc.add_paragraph('Backend Layer: Logic xử lý (Python, Flask, Scikit-learn)', style='List Bullet')
doc.add_paragraph('Database Layer: Lưu trữ dữ liệu (CSV, JSON, Pandas)', style='List Bullet')

# Section 2
doc.add_heading('2. Frontend Layer (Tầng Giao Diện)', 1)
doc.add_paragraph('Frontend Layer là phần giao diện mà người dùng tương tác trực tiếp.')

doc.add_heading('2.1 Công Nghệ Sử Dụng', 2)
doc.add_paragraph('HTML (templates/index.html): Cấu trúc form, layout', style='List Bullet')
doc.add_paragraph('CSS (static/style.css): Styling, responsive design', style='List Bullet')
doc.add_paragraph('JavaScript (static/script.js): Tương tác, validation, API calls', style='List Bullet')

doc.add_heading('2.2 Chức Năng Chính', 2)
doc.add_paragraph('Bước 1: Hiển thị form nhập liệu', style='List Number')
doc.add_paragraph('8 trường: 6 categorical + 2 text', style='List Bullet')
doc.add_paragraph('Bước 2: Chuẩn hóa dữ liệu', style='List Number')
doc.add_paragraph('Remove diacritics (ă, ơ, ư, etc.)', style='List Bullet')
doc.add_paragraph('Chuyển sang không dấu, lowercase', style='List Bullet')
doc.add_paragraph('Bước 3: Gửi POST request đến /predict endpoint', style='List Number')
doc.add_paragraph('Bước 4: Nhận JSON response từ backend', style='List Number')
doc.add_paragraph('Bước 5: Render kết quả Top 3 ngành', style='List Number')
doc.add_paragraph('Hiển thị score, confidence, suggestions', style='List Bullet')

doc.add_heading('2.3 Form Input', 2)
table1 = doc.add_table(rows=4, cols=2)
table1.style = 'Light Grid Accent 1'
table1.rows[0].cells[0].text = 'Loại'
table1.rows[0].cells[1].text = 'Trường'
table1.rows[1].cells[0].text = 'Select (6)'
table1.rows[1].cells[1].text = 'interest, favorite_subject, personality, skills, work_environment, career_goal'
table1.rows[2].cells[0].text = 'Text Input (2)'
table1.rows[2].cells[1].text = 'self_description, future_direction'
table1.rows[3].cells[0].text = 'Validation'
table1.rows[3].cells[1].text = 'JavaScript: required field check, length check'

# Section 3
doc.add_heading('3. Backend Layer (Tầng Xử Lý)', 1)
doc.add_paragraph('Backend Layer là nơi xử lý logic của ứng dụng.')

doc.add_heading('3.1 Công Nghệ Sử Dụng', 2)
doc.add_paragraph('Python 3.x: Ngôn ngữ lập trình', style='List Bullet')
doc.add_paragraph('Flask: Web framework nhẹ', style='List Bullet')
doc.add_paragraph('Scikit-learn: Machine Learning library', style='List Bullet')
doc.add_paragraph('Joblib: Lưu/tải model (pickle)', style='List Bullet')
doc.add_paragraph('Pandas: Data manipulation', style='List Bullet')

doc.add_heading('3.2 Cấu Trúc Thư Mục', 2)
doc.add_paragraph('app.py: Flask main app, endpoints', style='List Bullet')
doc.add_paragraph('utils/predictor.py: Hybrid Model logic', style='List Bullet')
doc.add_paragraph('utils/features.py: Feature extraction', style='List Bullet')
doc.add_paragraph('utils/constants.py: Config constants', style='List Bullet')
doc.add_paragraph('utils/chatbot.py: Chatbot logic', style='List Bullet')

doc.add_heading('3.3 Endpoints', 2)
table2 = doc.add_table(rows=4, cols=3)
table2.style = 'Light Grid Accent 1'
table2.rows[0].cells[0].text = 'Method'
table2.rows[0].cells[1].text = 'Endpoint'
table2.rows[0].cells[2].text = 'Mô Tả'
table2.rows[1].cells[0].text = 'GET'
table2.rows[1].cells[1].text = '/'
table2.rows[1].cells[2].text = 'Trang chủ form'
table2.rows[2].cells[0].text = 'POST'
table2.rows[2].cells[1].text = '/predict'
table2.rows[2].cells[2].text = 'Nhận hồ sơ → trả Top 3'
table2.rows[3].cells[0].text = 'GET'
table2.rows[3].cells[1].text = '/health'
table2.rows[3].cells[2].text = 'Kiểm tra model status'

doc.add_heading('3.4 /predict Endpoint Workflow', 2)
doc.add_paragraph('Bước 1: Nhận JSON payload từ frontend', style='List Number')
doc.add_paragraph('Bước 2: Validate dữ liệu (kiểm tra required fields)', style='List Number')
doc.add_paragraph('Bước 3: Gọi Predictor.predict(payload)', style='List Number')
doc.add_paragraph('Bước 4: Predictor trả về dictionary với top_3 + scores', style='List Number')
doc.add_paragraph('Bước 5: Format response với major names + suggestions', style='List Number')
doc.add_paragraph('Bước 6: Return JSON response', style='List Number')

# Section 4
doc.add_heading('4. Database Layer (Tầng Dữ Liệu)', 1)
doc.add_paragraph('Database Layer lưu trữ tất cả dữ liệu của hệ thống.')

doc.add_heading('4.1 File CSV (Training Data)', 2)
doc.add_paragraph('File: data/raw/students.csv')
doc.add_paragraph('Chứa ~1200 mẫu học sinh')
doc.add_paragraph('Cột: 8 trường input + 1 cột target ngành')
doc.add_paragraph('Được tải lên khi training model')

doc.add_heading('4.2 File JSON (Cấu Hình & Dữ Liệu)', 2)
doc.add_paragraph('models/majors.json: 50+ ngành + mô tả', style='List Bullet')
doc.add_paragraph('models/hybrid_config.json: Weights (40% & 60%)', style='List Bullet')
doc.add_paragraph('data/majors_profiles.json: Profile descriptions', style='List Bullet')
doc.add_paragraph('data/salary_benchmarks.json: Salary info', style='List Bullet')

doc.add_heading('4.3 Pickle Files (Trained Models)', 2)
doc.add_paragraph('models/rf_model.pkl: RandomForest model đã train', style='List Bullet')
doc.add_paragraph('models/ohe.pkl: OneHotEncoder đã fit', style='List Bullet')
doc.add_paragraph('models/tfidf.pkl: TfidfVectorizer đã fit', style='List Bullet')
doc.add_paragraph('models/classes.pkl: List class labels', style='List Bullet')

doc.add_heading('4.4 Data Loading Process', 2)
doc.add_paragraph('Bước 1: app.py khởi động', style='List Number')
doc.add_paragraph('Bước 2: Load pickle files từ models/', style='List Number')
doc.add_paragraph('Bước 3: Load JSON files từ data/', style='List Number')
doc.add_paragraph('Bước 4: Initialize Predictor với loaded models', style='List Number')
doc.add_paragraph('Bước 5: Predictor sẵn sàng xử lý requests', style='List Number')

# Section 5
doc.add_heading('5. Data Flow (Luồng Dữ Liệu)', 1)

doc.add_heading('5.1 Request Flow', 2)
doc.add_paragraph('User fills form → Frontend chuẩn hóa → POST /predict → Backend process → JSON response → Frontend render')

doc.add_heading('5.2 Chi Tiết Data Flow', 2)
doc.add_paragraph('Bước 1: User Input', style='List Number')
doc.add_paragraph('User chọn categorical values + nhập text', style='List Bullet')
doc.add_paragraph('Bước 2: Frontend Processing', style='List Number')
doc.add_paragraph('JavaScript chuẩn hóa: remove diacritics', style='List Bullet')
doc.add_paragraph('Tạo JSON payload', style='List Bullet')
doc.add_paragraph('POST request đến http://localhost:5000/predict', style='List Bullet')
doc.add_paragraph('Bước 3: Backend Receives', style='List Number')
doc.add_paragraph('Flask route /predict nhận request', style='List Bullet')
doc.add_paragraph('Parse JSON body', style='List Bullet')
doc.add_paragraph('Bước 4: Hybrid Processing', style='List Number')
doc.add_paragraph('Content-Based: TF-IDF + Cosine Similarity', style='List Bullet')
doc.add_paragraph('Machine Learning: One-Hot + RandomForest', style='List Bullet')
doc.add_paragraph('Combination: 0.40 × Content + 0.60 × ML', style='List Bullet')
doc.add_paragraph('Ranking & Top 3 selection', style='List Bullet')
doc.add_paragraph('Bước 5: Response Generation', style='List Number')
doc.add_paragraph('Format top_3 + scores + suggestions', style='List Bullet')
doc.add_paragraph('Return JSON response', style='List Bullet')
doc.add_paragraph('Bước 6: Frontend Rendering', style='List Number')
doc.add_paragraph('Parse JSON response', style='List Bullet')
doc.add_paragraph('Render Top 3 ngành with UI components', style='List Bullet')

doc.add_page_break()

# Section 6
doc.add_heading('6. Technology Stack Summary', 1)
table3 = doc.add_table(rows=5, cols=3)
table3.style = 'Light Grid Accent 1'
table3.rows[0].cells[0].text = 'Layer'
table3.rows[0].cells[1].text = 'Technology'
table3.rows[0].cells[2].text = 'Purpose'
table3.rows[1].cells[0].text = 'Frontend'
table3.rows[1].cells[1].text = 'HTML5, CSS3, JavaScript'
table3.rows[1].cells[2].text = 'User interface & interaction'
table3.rows[2].cells[0].text = 'Backend'
table3.rows[2].cells[1].text = 'Python, Flask, Scikit-learn'
table3.rows[2].cells[2].text = 'Business logic & prediction'
table3.rows[3].cells[0].text = 'Data'
table3.rows[3].cells[1].text = 'CSV, JSON, Pickle'
table3.rows[3].cells[2].text = 'Storage & persistence'
table3.rows[4].cells[0].text = 'ML'
table3.rows[4].cells[1].text = 'RandomForest, TF-IDF, OHE'
table3.rows[4].cells[2].text = 'Prediction & scoring'

# Section 7
doc.add_heading('7. Deployment', 1)
doc.add_paragraph('Bước 1: Setup Environment', style='List Number')
doc.add_paragraph('python -m venv venv', style='List Bullet')
doc.add_paragraph('source venv/bin/activate (hoặc venv\\Scripts\\activate trên Windows)', style='List Bullet')
doc.add_paragraph('Bước 2: Install Dependencies', style='List Number')
doc.add_paragraph('pip install -r requirements.txt', style='List Bullet')
doc.add_paragraph('Bước 3: Train Model (nếu cần)', style='List Number')
doc.add_paragraph('python train_model.py', style='List Bullet')
doc.add_paragraph('Bước 4: Run Application', style='List Number')
doc.add_paragraph('python app.py', style='List Bullet')
doc.add_paragraph('Bước 5: Access Application', style='List Number')
doc.add_paragraph('Open browser → http://127.0.0.1:5000', style='List Bullet')

# Section 8 - Conclusion
doc.add_heading('8. Kết Luận', 1)
doc.add_paragraph('Kiến trúc 3-Tier được sử dụng để:')
doc.add_paragraph('✓ Tách biệt concerns (Frontend, Backend, Database)', style='List Bullet')
doc.add_paragraph('✓ Dễ bảo trì và mở rộng', style='List Bullet')
doc.add_paragraph('✓ Cho phép thay đổi từng layer độc lập', style='List Bullet')
doc.add_paragraph('✓ Cải thiện performance thông qua caching', style='List Bullet')
doc.add_paragraph('✓ Hỗ trợ scalability cho tương lai', style='List Bullet')

doc.save('ARCHITECTURE_MODULE.docx')
print("✅ Architecture Module completed: ARCHITECTURE_MODULE.docx")
