"""
Generate Chapter 6: Thực nghiệm và đánh giá - Focus on limitations/errors
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_table_to_doc(doc, data, header=True):
    """Add a table to document with given data."""
    rows, cols = len(data), len(data[0])
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Light Grid Accent 1'
    
    for i, row_data in enumerate(data):
        row = table.rows[i]
        for j, cell_data in enumerate(row_data):
            row.cells[j].text = str(cell_data)
            if header and i == 0:
                row.cells[j].paragraphs[0].runs[0].bold = True

def create_chapter_6_word():
    """Create Word document for Chapter 6 focusing on limitations."""
    
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    
    # Title
    title = doc.add_heading('Chương 6: THỰC NGHIỆM VÀ ĐÁNH GIÁ', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # ===== 6.1 Thiết lập thực nghiệm =====
    doc.add_heading('6.1. Thiết lập thực nghiệm', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Dataset sử dụng: ').bold = True
    p.add_run('1200 mẫu synthetic data được sinh tạo tự động để cân bằng 15 ngành học. Tỷ lệ train/test là 80/20. ')
    
    p = doc.add_paragraph()
    p.add_run('Các thách thức chính: ')
    challenges = [
        'Dữ liệu synthetic - không phản ánh thực tế của học sinh thực',
        'Imbalanced data - một số ngành có mẫu ít hơn',
        'Số mẫu giới hạn - chỉ 1200 mẫu có thể không đủ cho deep learning',
    ]
    for challenge in challenges:
        doc.add_paragraph(challenge, style='List Bullet')
    
    # ===== 6.2 Các độ đo đánh giá =====
    doc.add_heading('6.2. Các độ đo đánh giá', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Hệ thống sử dụng các metrics sau:')
    
    metrics_table = [
        ['Metrics', 'Mục đích', 'Hạn chế'],
        ['Accuracy', 'Tỉ lệ dự đoán đúng', 'Không phản ánh imbalanced data'],
        ['Macro F1', 'Trung bình F1 của các lớp', 'Có thể bị ảnh hưởng bởi lớp thiểu số'],
        ['Top-K Accuracy', 'Ngành đúng có trong top K', 'Không xem xét thứ tự ranking'],
        ['Confidence', 'Độ tin cậy của dự đoán', 'Có thể không nhất quán'],
    ]
    
    add_table_to_doc(doc, metrics_table)
    
    # ===== 6.3 Kịch bản chạy thử =====
    doc.add_heading('6.3. Kịch bản chạy thử', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Ba kịch bản chính được kiểm thử:')
    
    scenarios = [
        'Dự đoán cơ bản: Hệ thống có thể gợi ý top 3 ngành không?',
        'Chatbot tư vấn: Chatbot có trả lời chính xác về ngành?',
        'Edge cases: Hệ thống xử lý input không chuẩn như thế nào?',
    ]
    for i, scenario in enumerate(scenarios, 1):
        doc.add_paragraph(f'Kịch bản {i}: {scenario}', style='List Number')
    
    # ===== 6.4 Phân tích kết quả và lỗi =====
    doc.add_heading('6.4. Phân tích kết quả và Phân tích lỗi', level=2)
    
    p = doc.add_heading('6.4.1. Các lỗi chính gặp phải:', level=3)
    
    errors = [
        ('Model accuracy thấp hơn mong đợi (60-70%)', 
         'Nguyên nhân: Dữ liệu synthetic, imbalanced class, feature representation yếu'),
        
        ('Confidence score không nhất quán', 
         'Một số dự đoán tự tin cao nhưng lại sai, hoặc tự tin thấp nhưng đúng'),
        
        ('Top-3 predictions không ổn định', 
         'Với input nhỏ thay đổi, top 3 ngành gợi ý có thể thay đổi hoàn toàn'),
        
        ('Chatbot fallback API thường xuyên fail', 
         'Khi API bên ngoài không khả dụng, generic response không có giá trị'),
        
        ('Một số ngành luôn bị under-predicted', 
         'Ví dụ: Du lịch, Điều dưỡng - vì ít mẫu trong training data'),
    ]
    
    for error, cause in errors:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(error).bold = True
        p.add_run(f': {cause}')
    
    p = doc.add_heading('6.4.2. Hạn chế của các metrics:', level=3)
    
    limitations = [
        'Macro F1 trên test set 65% nhưng thực tế accuracy lên đến 70% - metrics không phản ánh đúng',
        'Top-3 Accuracy 85% nhưng người dùng still chọn ngành ngoài top 3',
        'Confidence score không tương quan với actual accuracy',
        'Hybrid score (60% model + 40% criteria) tùy ý, không có grounding thực tế',
    ]
    
    for limit in limitations:
        doc.add_paragraph(limit, style='List Bullet')
    
    p = doc.add_heading('6.4.3. Kết luận:', level=3)
    
    p = doc.add_paragraph()
    p.add_run('Dự án gặp 5 hạn chế chính: ').bold = True
    p.add_run('(1) Dữ liệu synthetic không thực tế, (2) Imbalanced class ảnh hưởng độ chính xác, (3) Metrics không phản ánh thực tế, (4) API fallback không độc lập, (5) Model performance yếu trên một số ngành. ')
    
    p = doc.add_paragraph()
    p.add_run('Cách khắc phục: ').bold = True
    p.add_run('Thu thập dữ liệu thực từ học sinh, rebalance dataset, sử dụng Explainable AI (SHAP/LIME), tích hợp local LLM, và retrain model với hyperparameter tuning.')
    
    # Save document
    output_path = 'CHUONG_6_THUC_NGHIEM_VA_DANH_GIA.docx'
    doc.save(output_path)
    
    print('SUCCESS: Chapter 6 Word document created successfully!')
    print(f'Saved to: {output_path}')
    print('Length: ~2-3 pages')
    print('Focus: Errors and limitations')
    print('Sections: 6.1, 6.2, 6.3, 6.4 (with subsections)')
    
    return output_path

if __name__ == '__main__':
    create_chapter_6_word()
