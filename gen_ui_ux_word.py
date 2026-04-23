from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('MODULE 5: GIAO DIỆN NGƯỜI DÙNG (USER INTERFACE / UX)', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Introduction
doc.add_heading('1. Giới Thiệu', 1)
doc.add_paragraph('Module UI/UX mô tả giao diện người dùng của hệ thống tư vấn ngành học.')
doc.add_paragraph('Giao diện được thiết kế để:')
doc.add_paragraph('Dễ sử dụng (User-Friendly)', style='List Bullet')
doc.add_paragraph('Minh bạch và rõ ràng', style='List Bullet')
doc.add_paragraph('Responsive (hoạt động trên mobile & desktop)', style='List Bullet')
doc.add_paragraph('Hỗ trợ tiếng Việt đầy đủ', style='List Bullet')

# Section 2
doc.add_heading('2. Các Thành Phần Giao Diện', 1)

doc.add_heading('2.1 5.3.1: Web Form Input (Biểu Mẫu Nhập Liệu)', 2)
doc.add_paragraph('Phần form cho người dùng nhập thông tin về bản thân.')

doc.add_heading('2.1.1 Bố Cục Form', 2)
doc.add_paragraph('Header: Tiêu đề "Tư Vấn Ngành Học" + hướng dẫn', style='List Bullet')
doc.add_paragraph('Section 1: Sở thích chính (Select dropdown)', style='List Bullet')
doc.add_paragraph('Section 2: Môn học yêu thích (Select dropdown)', style='List Bullet')
doc.add_paragraph('Section 3: Kỹ năng nổi bật (Select dropdown)', style='List Bullet')
doc.add_paragraph('Section 4: Tính cách (Select dropdown)', style='List Bullet')
doc.add_paragraph('Section 5: Môi trường làm việc mong muốn (Select dropdown)', style='List Bullet')
doc.add_paragraph('Section 6: Mục tiêu nghề nghiệp (Select dropdown)', style='List Bullet')
doc.add_paragraph('Section 7: Mô tả bản thân (Text Area)', style='List Bullet')
doc.add_paragraph('Section 8: Định hướng tương lai (Text Area)', style='List Bullet')

doc.add_heading('2.1.2 Styling & Màu Sắc', 2)
doc.add_paragraph('Input fields: Trắng background, border xám nhẹ', style='List Bullet')
doc.add_paragraph('Labels: Đen, font size 14px, bold', style='List Bullet')
doc.add_paragraph('Text areas: Cao 100px, placeholder text màu xám', style='List Bullet')
doc.add_paragraph('Focus state: Border xanh dương, shadow nhẹ', style='List Bullet')

doc.add_heading('2.1.3 Validation', 2)
doc.add_paragraph('All 6 select fields: Required (bắt buộc)', style='List Bullet')
doc.add_paragraph('2 text areas: Optional (tùy chọn)', style='List Bullet')
doc.add_paragraph('Text area max length: 500 ký tự', style='List Bullet')
doc.add_paragraph('Submit button: Disabled nếu form chưa hoàn thành', style='List Bullet')

# Section 3
doc.add_heading('2.2 5.3.2: Nút Dự Đoán & Color Coding', 2)

doc.add_heading('2.2.1 Nút Dự Đoán (Predict Button)', 2)
doc.add_paragraph('Text: "🚀 Dự Đoán Ngành Cho Bạn"', style='List Bullet')
doc.add_paragraph('Background: Gradient xanh dương → xanh lá', style='List Bullet')
doc.add_paragraph('Text color: Trắng, font-weight bold', style='List Bullet')
doc.add_paragraph('Size: 16px, padding 12px 24px', style='List Bullet')
doc.add_paragraph('Hover state: Màu đậm hơn, shadow, cursor pointer', style='List Bullet')
doc.add_paragraph('Click state: Loading animation (spinner)', style='List Bullet')

doc.add_heading('2.2.2 Color Coding (Mã Màu Sắc)', 2)
doc.add_paragraph('Xanh dương (#2196F3): Primary action, info', style='List Bullet')
doc.add_paragraph('Xanh lá (#4CAF50): Success, confidence cao', style='List Bullet')
doc.add_paragraph('Cam (#FF9800): Warning, confidence trung bình', style='List Bullet')
doc.add_paragraph('Đỏ (#F44336): Error, confidence thấp', style='List Bullet')
doc.add_paragraph('Xám (#9E9E9E): Disabled, neutral', style='List Bullet')

# Section 4
doc.add_heading('2.3 5.3.3: Giao Diện Hiển Thị Kết Quả', 2)

doc.add_heading('2.3.1 Cấu Trúc Hiển Thị Top 3', 2)
doc.add_paragraph('Bước 1: Hiển thị loading animation (500ms)', style='List Number')
doc.add_paragraph('Bước 2: Animate results in từ dưới lên', style='List Number')
doc.add_paragraph('Bước 3: Hiển thị Top 3 cards (từng cái một)', style='List Number')

doc.add_heading('2.3.2 Card Layout cho Mỗi Ngành', 2)
doc.add_paragraph('Rank badge: #1, #2, #3 (icon + số)', style='List Bullet')
doc.add_paragraph('Major name: Font size 18px, bold', style='List Bullet')
doc.add_paragraph('Score: "Điểm: XX/100" với color coding', style='List Bullet')
doc.add_paragraph('Confidence: "Độ tin cậy: XX/100 (Cao/Trung bình/Tham khảo)"', style='List Bullet')
doc.add_paragraph('Gap from next: "Chênh ngành kế tiếp: +X.XX điểm"', style='List Bullet')
doc.add_paragraph('Feedback: "Bạn có mức phù hợp cao với ngành này"', style='List Bullet')
doc.add_paragraph('Button: "Xem Chi Tiết" (link to major info)', style='List Bullet')

doc.add_heading('2.3.3 Score Visualization', 2)
doc.add_paragraph('Progress bar: Chiều dài tương ứng điểm (0-100)', style='List Bullet')
doc.add_paragraph('Bar color: Xanh (>70), Cam (50-70), Đỏ (<50)', style='List Bullet')
doc.add_paragraph('Smooth animation: 1 second transition', style='List Bullet')

doc.add_heading('2.3.4 Responsive Design', 2)
doc.add_paragraph('Desktop: 3 cards ngang (1 card = 30% width)', style='List Bullet')
doc.add_paragraph('Tablet: 2 cards ngang (1 card = 45% width)', style='List Bullet')
doc.add_paragraph('Mobile: 1 card dọc (1 card = 100% width)', style='List Bullet')
doc.add_paragraph('Padding: Auto adjust dựa screen size', style='List Bullet')

# Section 5
doc.add_heading('2.4 5.3.4: Nút Chatbot & Giao Diện Chat', 2)

doc.add_heading('2.4.1 Chatbot Button', 2)
doc.add_paragraph('Position: Fixed, bottom-right corner', style='List Bullet')
doc.add_paragraph('Icon: Chat bubble 💬', style='List Bullet')
doc.add_paragraph('Background: Xanh dương (#2196F3)', style='List Bullet')
doc.add_paragraph('Size: 60px × 60px (rounded)', style='List Bullet')
doc.add_paragraph('Shadow: Drop shadow 4px', style='List Bullet')
doc.add_paragraph('Hover: Pulse animation', style='List Bullet')

doc.add_heading('2.4.2 Chat Window', 2)
doc.add_paragraph('Position: Pop-up từ bottom-right', style='List Bullet')
doc.add_paragraph('Size: 350px × 500px (responsive)', style='List Bullet')
doc.add_paragraph('Header: "AI Tư Vấn Ngành Học" + close button (X)', style='List Bullet')
doc.add_paragraph('Messages area: Scrollable, max height 400px', style='List Bullet')
doc.add_paragraph('Input field: Text input + send button', style='List Bullet')
doc.add_paragraph('Border radius: 12px (rounded corners)', style='List Bullet')

doc.add_heading('2.4.3 Chat Message Styling', 2)
doc.add_paragraph('User message: Xanh dương, bubble nền, align right', style='List Bullet')
doc.add_paragraph('Bot message: Xám, bubble nền, align left', style='List Bullet')
doc.add_paragraph('Timestamp: Font size 12px, màu xám', style='List Bullet')
doc.add_paragraph('Animation: Slide in từ left/right', style='List Bullet')

# Section 6
doc.add_heading('3. Tương Tác (Interactions)', 1)

doc.add_heading('3.1 Form Submission Flow', 2)
doc.add_paragraph('User điền form đầy đủ', style='List Number')
doc.add_paragraph('Click "Dự Đoán Ngành"', style='List Number')
doc.add_paragraph('Button hiển thị loading state', style='List Number')
doc.add_paragraph('Frontend gửi POST /predict', style='List Number')
doc.add_paragraph('Backend xử lý (1-2 giây)', style='List Number')
doc.add_paragraph('Results animate in', style='List Number')
doc.add_paragraph('User có thể scroll down để xem Full details', style='List Number')

doc.add_heading('3.2 Chat Interaction Flow', 2)
doc.add_paragraph('Click chatbot icon', style='List Number')
doc.add_paragraph('Chat window mở với greeting message', style='List Number')
doc.add_paragraph('User nhập câu hỏi', style='List Number')
doc.add_paragraph('Click send hoặc nhấn Enter', style='List Number')
doc.add_paragraph('Bot xử lý và trả lời (1-3 giây)', style='List Number')
doc.add_paragraph('Message hiển thị trong chat window', style='List Number')

# Section 7
doc.add_heading('4. Accessibility (Khả Năng Tiếp Cận)', 1)

doc.add_paragraph('Alt text cho tất cả images', style='List Bullet')
doc.add_paragraph('Keyboard navigation (Tab, Enter)', style='List Bullet')
doc.add_paragraph('Color contrast: WCAG AA standard', style='List Bullet')
doc.add_paragraph('Font size readable: Min 14px', style='List Bullet')
doc.add_paragraph('Mobile touch-friendly: Min 44px buttons', style='List Bullet')
doc.add_paragraph('Screen reader support: Proper semantic HTML', style='List Bullet')

# Section 8 - Conclusion
doc.add_heading('5. Kết Luận', 1)
doc.add_paragraph('UI/UX được thiết kế theo nguyên tắc:')
doc.add_paragraph('✓ Simplicity: Giao diện đơn giản, không phức tạp', style='List Bullet')
doc.add_paragraph('✓ Clarity: Rõ ràng, dễ hiểu', style='List Bullet')
doc.add_paragraph('✓ Consistency: Consistent styling & interactions', style='List Bullet')
doc.add_paragraph('✓ Responsiveness: Hoạt động tốt trên mọi thiết bị', style='List Bullet')
doc.add_paragraph('✓ Accessibility: Dễ tiếp cận cho tất cả người dùng', style='List Bullet')

doc.save('UI_UX_MODULE.docx')
print("✅ UI/UX Module completed: UI_UX_MODULE.docx")
