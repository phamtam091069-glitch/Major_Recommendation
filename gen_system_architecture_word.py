from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
title = doc.add_heading('MODULE 5: KIẾN TRÚC HỆ THỐNG (SYSTEM ARCHITECTURE)', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('1. Giới Thiệu', 1)
doc.add_paragraph('Module System Architecture mô tả cấu trúc tổng thể hệ thống.')
doc.add_paragraph('Gồm 3 phần chính: Kiến trúc tổng quát + Chức năng + Luồng dữ liệu', style='List Bullet')

doc.add_heading('2. 5.1.1: Sơ Đồ Hệ Thống Tổng Quát (3-Tier Architecture)', 2)

doc.add_heading('2.1 Mô Hình 3-Tier Architecture', 2)
doc.add_paragraph('Tầng 1: Client (Presentation Layer)', style='List Bullet')
doc.add_paragraph('Frontend HTML/CSS/JavaScript', style='List Bullet 2')
doc.add_paragraph('User Interface (Form + Results)', style='List Bullet 2')
doc.add_paragraph('Tầng 2: Server (Application Layer)', style='List Bullet')
doc.add_paragraph('Flask Backend API', style='List Bullet 2')
doc.add_paragraph('Business Logic (Prediction)', style='List Bullet 2')
doc.add_paragraph('Tầng 3: Local Storage (Data Layer)', style='List Bullet')
doc.add_paragraph('Models folder (pkl files)', style='List Bullet 2')
doc.add_paragraph('Data folder (CSV, JSON)', style='List Bullet 2')

doc.add_heading('2.2 Tầng 1: Client Side (Frontend)', 2)
doc.add_paragraph('Thành phần:', style='List Bullet')
doc.add_paragraph('templates/index.html: Form & Results display', style='List Bullet 2')
doc.add_paragraph('static/script.js: Event handling & API calls', style='List Bullet 2')
doc.add_paragraph('static/style.css: Styling & Responsive design', style='List Bullet 2')
doc.add_paragraph('Chức năng:', style='List Bullet')
doc.add_paragraph('Nhận input từ user', style='List Bullet 2')
doc.add_paragraph('Chuẩn hóa dữ liệu (remove diacritics)', style='List Bullet 2')
doc.add_paragraph('Gửi POST request tới /predict', style='List Bullet 2')
doc.add_paragraph('Hiển thị kết quả (Top 3 majors)', style='List Bullet 2')
doc.add_paragraph('Xử lý errors & loading states', style='List Bullet 2')

doc.add_heading('2.3 Tầng 2: Server Side (Backend)', 2)
doc.add_paragraph('Thành phần:', style='List Bullet')
doc.add_paragraph('app.py: Flask application & routes', style='List Bullet 2')
doc.add_paragraph('utils/predictor.py: Prediction logic', style='List Bullet 2')
doc.add_paragraph('utils/chatbot.py: Chatbot logic', style='List Bullet 2')
doc.add_paragraph('utils/constants.py: Constants & configurations', style='List Bullet 2')
doc.add_paragraph('Chức năng:', style='List Bullet')
doc.add_paragraph('GET /: Serve form page', style='List Bullet 2')
doc.add_paragraph('POST /predict: Process prediction', style='List Bullet 2')
doc.add_paragraph('GET /health: Check status', style='List Bullet 2')
doc.add_paragraph('POST /chat: Chatbot interaction', style='List Bullet 2')

doc.add_heading('2.4 Tầng 3: Storage Layer (Local)', 2)
doc.add_paragraph('Models folder:', style='List Bullet')
doc.add_paragraph('rf_model.pkl: Trained RandomForest', style='List Bullet 2')
doc.add_paragraph('ohe.pkl: OneHotEncoder', style='List Bullet 2')
doc.add_paragraph('tfidf.pkl: TfidfVectorizer', style='List Bullet 2')
doc.add_paragraph('classes.pkl: Class labels', style='List Bullet 2')
doc.add_paragraph('majors.json: Major descriptions', style='List Bullet 2')
doc.add_paragraph('Data folder:', style='List Bullet')
doc.add_paragraph('raw/students.csv: Training data', style='List Bullet 2')
doc.add_paragraph('majors_profiles.json: Major info', style='List Bullet 2')
doc.add_paragraph('salary_benchmarks.json: Salary data', style='List Bullet 2')

doc.add_heading('2.5 3-Tier Diagram Text', 2)
doc.add_paragraph('┌─────────────────────────────────┐', style='List Bullet')
doc.add_paragraph('│  TIER 1: CLIENT (Frontend)      │', style='List Bullet')
doc.add_paragraph('│  HTML/CSS/JavaScript            │', style='List Bullet')
doc.add_paragraph('│  Form + Results Display         │', style='List Bullet')
doc.add_paragraph('└────────────────┬────────────────┘', style='List Bullet')
doc.add_paragraph('                 │ HTTP/REST API', style='List Bullet')
doc.add_paragraph('┌────────────────▼────────────────┐', style='List Bullet')
doc.add_paragraph('│  TIER 2: SERVER (Backend)       │', style='List Bullet')
doc.add_paragraph('│  Flask + Python                 │', style='List Bullet')
doc.add_paragraph('│  /predict, /chat, /health      │', style='List Bullet')
doc.add_paragraph('└────────────────┬────────────────┘', style='List Bullet')
doc.add_paragraph('                 │ File I/O', style='List Bullet')
doc.add_paragraph('┌────────────────▼────────────────┐', style='List Bullet')
doc.add_paragraph('│  TIER 3: STORAGE (Local)        │', style='List Bullet')
doc.add_paragraph('│  Models (pkl) + Data (CSV/JSON) │', style='List Bullet')
doc.add_paragraph('│  Majors Info + Salary Data      │', style='List Bullet')
doc.add_paragraph('└─────────────────────────────────┘', style='List Bullet')

doc.add_heading('3. 5.1.2: Sơ Đồ Chức Năng của Hệ Thống', 2)

doc.add_heading('3.1 Các Chức Năng Chính', 2)
doc.add_paragraph('1. Prediction Module (Dự Đoán Ngành):', style='List Bullet')
doc.add_paragraph('Input: User profile (8 fields)', style='List Bullet 2')
doc.add_paragraph('Process: ML + Content-based scoring', style='List Bullet 2')
doc.add_paragraph('Output: Top 3 majors + scores', style='List Bullet 2')
doc.add_paragraph('2. Chatbot Module (Chat AI):', style='List Bullet')
doc.add_paragraph('Input: User questions', style='List Bullet 2')
doc.add_paragraph('Process: NLP + Pattern matching', style='List Bullet 2')
doc.add_paragraph('Output: Contextual responses', style='List Bullet 2')
doc.add_paragraph('3. Data Management (Quản Lý Dữ Liệu):', style='List Bullet')
doc.add_paragraph('Load models & data at startup', style='List Bullet 2')
doc.add_paragraph('Cache frequently used data', style='List Bullet 2')

doc.add_heading('3.2 Feature Modules', 2)
doc.add_paragraph('OneHot Encoder:', style='List Bullet')
doc.add_paragraph('Convert categorical to binary vectors', style='List Bullet 2')
doc.add_paragraph('TF-IDF Vectorizer:', style='List Bullet')
doc.add_paragraph('Convert text to weighted term vectors', style='List Bullet 2')
doc.add_paragraph('RandomForest Classifier:', style='List Bullet')
doc.add_paragraph('Multi-class classification', style='List Bullet 2')
doc.add_paragraph('Predictor:', style='List Bullet')
doc.add_paragraph('Combine all components', style='List Bullet 2')

doc.add_heading('3.3 Data Flow Diagram', 2)
doc.add_paragraph('User Input', style='List Bullet')
doc.add_paragraph('↓', style='List Bullet')
doc.add_paragraph('Frontend Validation & Normalization', style='List Bullet')
doc.add_paragraph('↓', style='List Bullet')
doc.add_paragraph('POST /predict (JSON payload)', style='List Bullet')
doc.add_paragraph('↓', style='List Bullet')
doc.add_paragraph('Backend: OneHot + TF-IDF encoding', style='List Bullet')
doc.add_paragraph('↓', style='List Bullet')
doc.add_paragraph('RandomForest Prediction', style='List Bullet')
doc.add_paragraph('↓', style='List Bullet')
doc.add_paragraph('Scoring & Ranking (Top 3)', style='List Bullet')
doc.add_paragraph('↓', style='List Bullet')
doc.add_paragraph('JSON Response with results', style='List Bullet')
doc.add_paragraph('↓', style='List Bullet')
doc.add_paragraph('Frontend: Display & Animation', style='List Bullet')

doc.add_heading('4. 5.1.3: Sơ Đồ Luồng Dữ Liệu (API & File Exchange)', 2)

doc.add_heading('4.1 API Endpoints', 2)
table_api = doc.add_table(rows=5, cols=4)
table_api.style = 'Light Grid Accent 1'
table_api.rows[0].cells[0].text = 'Method'
table_api.rows[0].cells[1].text = 'Endpoint'
table_api.rows[0].cells[2].text = 'Request'
table_api.rows[0].cells[3].text = 'Response'
table_api.rows[1].cells[0].text = 'GET'
table_api.rows[1].cells[1].text = '/'
table_api.rows[1].cells[2].text = 'None'
table_api.rows[1].cells[3].text = 'HTML form'
table_api.rows[2].cells[0].text = 'POST'
table_api.rows[2].cells[1].text = '/predict'
table_api.rows[2].cells[2].text = 'JSON payload'
table_api.rows[2].cells[3].text = 'Top 3 + scores'
table_api.rows[3].cells[0].text = 'GET'
table_api.rows[3].cells[1].text = '/health'
table_api.rows[3].cells[2].text = 'None'
table_api.rows[3].cells[3].text = 'Status JSON'
table_api.rows[4].cells[0].text = 'POST'
table_api.rows[4].cells[1].text = '/chat'
table_api.rows[4].cells[2].text = 'Message JSON'
table_api.rows[4].cells[3].text = 'Bot response'

doc.add_heading('4.2 Request Payload Example (/predict)', 2)
doc.add_paragraph('{', style='List Bullet')
doc.add_paragraph('"so_thich_chinh": "cong nghe"', style='List Bullet 2')
doc.add_paragraph('"mon_hoc_yeu_thich": "tin hoc"', style='List Bullet 2')
doc.add_paragraph('"ky_nang_noi_bat": "lap trinh"', style='List Bullet 2')
doc.add_paragraph('"tinh_cach": "tho thao"', style='List Bullet 2')
doc.add_paragraph('"moi_truong_lam_viec": "cong ty"', style='List Bullet 2')
doc.add_paragraph('"muc_tieu_nghe_nghiep": "phat trien"', style='List Bullet 2')
doc.add_paragraph('"mo_ta_ban_than": "..."', style='List Bullet 2')
doc.add_paragraph('"dinh_huong_tuong_lai": "..."', style='List Bullet 2')
doc.add_paragraph('}', style='List Bullet')

doc.add_heading('4.3 Response Payload Example (/predict)', 2)
doc.add_paragraph('{', style='List Bullet')
doc.add_paragraph('"top_3": [', style='List Bullet 2')
doc.add_paragraph('{', style='List Bullet 2')
doc.add_paragraph('"major": "Cong nghe thong tin"', style='List Bullet 2')
doc.add_paragraph('"score": 85', style='List Bullet 2')
doc.add_paragraph('"confidence": "Cao"', style='List Bullet 2')
doc.add_paragraph('}', style='List Bullet 2')
doc.add_paragraph(']', style='List Bullet 2')
doc.add_paragraph('}', style='List Bullet')

doc.add_heading('4.4 File Exchange Flow', 2)
doc.add_paragraph('Startup:', style='List Bullet')
doc.add_paragraph('Load rf_model.pkl, ohe.pkl, tfidf.pkl', style='List Bullet 2')
doc.add_paragraph('Load majors.json, salary_benchmarks.json', style='List Bullet 2')
doc.add_paragraph('Runtime:', style='List Bullet')
doc.add_paragraph('Read from CSV if needed', style='List Bullet 2')
doc.add_paragraph('Write feedback to JSON (optional)', style='List Bullet 2')

doc.add_heading('5. Kết Luận', 1)
doc.add_paragraph('✓ 3-Tier architecture: Clear separation of concerns', style='List Bullet')
doc.add_paragraph('✓ Client: Handles UI & user input', style='List Bullet')
doc.add_paragraph('✓ Server: Processes predictions & logic', style='List Bullet')
doc.add_paragraph('✓ Storage: Manages models & data', style='List Bullet')
doc.add_paragraph('✓ REST API: Enables communication', style='List Bullet')
doc.add_paragraph('✓ File-based: Simple & scalable storage', style='List Bullet')

doc.save('SYSTEM_ARCHITECTURE_MODULE.docx')
print("✅ System Architecture Module completed: SYSTEM_ARCHITECTURE_MODULE.docx")
