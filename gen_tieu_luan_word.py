"""
Generate a concise thesis document highlighting limitations as a Word file.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_thesis_word():
    """Create a Word document with concise thesis focusing on limitations."""
    
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    
    # Title
    title = doc.add_heading('Chương 7: KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 7.1 Kết luận
    doc.add_heading('7.1. Kết luận', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Dự án thành công ').bold = True
    p.add_run('trong việc xây dựng hệ thống AI tư vấn ngành học toàn diện. Hệ thống kết hợp Machine Learning (Random Forest 30%) với tiêu chí minh bạch (70%), tích hợp chatbot AI và cơ chế fallback API đa tầng, cho phép chuẩn hóa dữ liệu tự động. Điểm nổi bật là tính giải thích được - người dùng hiểu rõ tại sao hệ thống gợi ý ngành nào. Hệ thống có giá trị thực tiễn cao cho học sinh, nhà trường, và nhà phát triển.')
    
    # 7.2 Hạn chế (MAIN FOCUS)
    doc.add_heading('7.2. Hạn chế', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Dự án tồn tại ').bold = False
    p.add_run('nhiều hạn chế cần khắc phục:').bold = True
    
    limitations = [
        ('Dữ liệu synthetic', 'Model được huấn luyện trên dữ liệu sinh tạo, không phản ánh thực tế. Dẫn đến overfitting, model không tổng quát hóa tốt với dữ liệu thực. Giải pháp: giảm trọng số model xuống 30%, tăng criteria score lên 70%.'),
        
        ('Dataset không đủ', 'Số lượng mẫu có thể không đủ để model học được tất cả pattern phức tạp. Hiệu năng model không ổn định với một số loại input, khó xử lý edge cases.'),
        
        ('Model Black Box', 'Random Forest khó giải thích chính xác tại sao model đưa ra quyết định cụ thể. Khó debug khi model cho kết quả không mong muốn.'),
        
        ('Trọng số cố định', 'Trọng số tiêu chí (sở thích 23%, định hướng 20%, v.v.) không thích ứng. Không phù hợp cho các nhóm học sinh khác nhau hay bối cảnh địa phương khác.'),
        
        ('Phụ thuộc API bên ngoài', 'Chatbot dựa vào Claude, OpenAI, Deepseek. Nếu tất cả fail → generic response. Có latency, chi phí API, phụ thuộc nhà cung cấp.'),
        
        ('Fallback API giới hạn', 'Nếu dữ liệu hoàn toàn sai lệch, fallback API khó xử lý. Không có cơ chế học từ những trường hợp normalize thất bại.'),
        
        ('Cache TTL cố định', 'Cache 1 giờ không thích ứng. Có thể lãng phí bộ nhớ hoặc trả về dữ liệu cũ.'),
        
        ('Phạm vi hạn chế', 'Chỉ hỗ trợ 15 ngành (không bao phủ toàn bộ hệ thống giáo dục Việt). Chỉ tiếng Việt, khó mở rộng.'),
        
        ('Confidence score không chính xác', 'Độ tin cậy dựa trên fit score + độ tách biệt có thể không phản ánh độ tin cậy thực tế.'),
        
        ('UI/UX cần cải thiện', 'Giao diện hiện tại functional nhưng cần design hiện đại hơn, animation, mobile responsiveness.'),
        
        ('Không có admin dashboard', 'Khó quản lý dữ liệu, model, monitor performance trong production, trigger retrain.'),
    ]
    
    for title, desc in limitations:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(title).bold = True
        p.add_run(': ' + desc)
    
    # Summary of limitations
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Tóm lại: ').bold = True
    p.add_run('Dự án có 12 hạn chế chính, tập trung vào dữ liệu (synthetic, không đủ), model (black box, overfitting), hệ thống (API dependencies), phạm vi (15 ngành, 1 ngôn ngữ), và triển khai (UI/UX, admin).')
    
    # 7.3 Hướng phát triển
    doc.add_heading('7.3. Hướng phát triển', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Để khắc phục các hạn chế:')
    
    doc.add_heading('Ngắn hạn (3-6 tháng):', level=3)
    short_term = [
        'Thu thập dữ liệu thực từ người dùng, loại bỏ synthetic data',
        'Redesign UI/UX với framework modern (React/Vue), thêm visualizations',
        'Tích hợp local LLM để giảm latency, chi phí API',
        'Expand: 20-30 ngành mới, thêm job prospects, salary info',
    ]
    for item in short_term:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Trung hạn (6-12 tháng):', level=3)
    mid_term = [
        'Thử model khác (XGBoost, LightGBM, Neural Networks)',
        'Explainable AI (SHAP, LIME) để giải thích quyết định model',
        'Adaptive weights: điều chỉnh trọng số dựa trên feedback',
        'Admin dashboard: quản lý ngành, dữ liệu, model, monitor metrics',
        'Multi-language support: Anh, Trung, Nhật',
    ]
    for item in mid_term:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Dài hạn (12+ tháng):', level=3)
    long_term = [
        'Scrape job postings, salary data, university rankings (real-time)',
        'Personalization engine: user profiling, collaborative filtering',
        'Integration: API for 3rd parties, SSO, webhooks',
        'Offline support: PWA, mobile app (iOS/Android)',
    ]
    for item in long_term:
        doc.add_paragraph(item, style='List Bullet')
    
    # Conclusion
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Dự án có tiềm năng trở thành platform hàng đầu nếu khắc phục được hạn chế về dữ liệu, model, và triển khai. Ưu tiên: (1) Thu thập dữ liệu thực, (2) Nâng cao độ chính xác model, (3) Cải thiện UX, (4) Mở rộng phạm vi.').italic = True
    
    # Save document
    output_path = 'TIEU_LUAN_KET_LUAN_HUONG_PHAT_TRIEN_NGAN_GON.docx'
    doc.save(output_path)
    
    print(f'✅ Word document created successfully!')
    print(f'📁 Saved to: {output_path}')
    print(f'📄 File type: DOCX (Microsoft Word)')
    print(f'📊 Content: Concise thesis focusing on 12 limitations')
    
    return output_path

if __name__ == '__main__':
    create_thesis_word()
