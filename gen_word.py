from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('MODULE CONTENT-BASED FILTERING', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Introduction
doc.add_heading('1. Giới Thiệu', 1)
doc.add_paragraph('Module Content-Based Filtering là một phương pháp trong hệ thống tư vấn ngành học, dựa trên việc so sánh nội dung (content) của hồ sơ học sinh với nội dung của các ngành học.')
doc.add_paragraph('Phương pháp này hoạt động bằng cách:')
doc.add_paragraph('Tạo "hồ sơ" cho mỗi học sinh dựa trên sở thích, kỹ năng và mục tiêu của họ', style='List Bullet')
doc.add_paragraph('Tạo "hồ sơ" cho mỗi ngành học dựa trên mô tả chi tiết về chương trình', style='List Bullet')
doc.add_paragraph('So sánh độ tương đồng giữa 2 hồ sơ bằng công thức toán học (Cosine Similarity)', style='List Bullet')

# Section 2
doc.add_heading('2. Hồ Sơ Ngành (Major Profile)', 1)
doc.add_paragraph('Hồ sơ ngành là biểu diễn chi tiết về một ngành học, được xây dựng từ thông tin có sẵn trong hệ thống.')

doc.add_heading('2.1 Nguồn Dữ Liệu', 2)
doc.add_paragraph('File: models/majors.json')
doc.add_paragraph('Chứa 50+ ngành học trong dự án')
doc.add_paragraph('Mỗi ngành có 2 trường chính:', style='List Bullet')
doc.add_paragraph('nganh: Tên ngành (được chuẩn hóa - không dấu, lowercase)', style='List Bullet 2')
doc.add_paragraph('mo_ta: Mô tả chi tiết ngành (text tự do)', style='List Bullet 2')

doc.add_heading('2.2 Ví Dụ Thực Tế', 2)
table1 = doc.add_table(rows=3, cols=2)
table1.style = 'Light Grid Accent 1'
table1.rows[0].cells[0].text = 'Ngành'
table1.rows[0].cells[1].text = 'Mô Tả'
table1.rows[1].cells[0].text = 'Cong nghe thong tin'
table1.rows[1].cells[1].text = 'Ngành Công nghệ thông tin đào tạo về lập trình, hệ thống máy tính, mạng, bảo mật...'
table1.rows[2].cells[0].text = 'Khoa hoc du lieu'
table1.rows[2].cells[1].text = 'Ngành Khoa học dữ liệu đào tạo cách thu thập, làm sạch, phân tích và trực quan hóa dữ liệu...'

doc.add_heading('2.3 Xử Lý Hồ Sơ Ngành', 2)
doc.add_paragraph('1. Tải từ JSON: Đọc file majors.json', style='List Number')
doc.add_paragraph('2. Chuẩn Hóa: Chuyển tên thành không dấu, lowercase', style='List Number')
doc.add_paragraph('3. Lưu Description: Mỗi ngành giữ mô tả nguyên gốc', style='List Number')
doc.add_paragraph('4. Sử Dụng cho TF-IDF: Các mô tả này sẽ được vectorize', style='List Number')

# Section 3
doc.add_heading('3. Hồ Sơ Học Sinh (Student Profile)', 1)
doc.add_paragraph('Hồ sơ học sinh là biểu diễn chi tiết về một học sinh, được xây dựng từ thông tin mà học sinh cung cấp.')

doc.add_heading('3.1 Các Thành Phần', 2)
doc.add_paragraph('Hồ sơ học sinh gồm 2 loại trường:')
doc.add_paragraph('Categorical Fields (6 trường chọn lựa):', style='List Bullet')
doc.add_paragraph('Sở thích chính (interest)', style='List Bullet 2')
doc.add_paragraph('Môn học yêu thích (favorite_subject)', style='List Bullet 2')
doc.add_paragraph('Tính cách (personality)', style='List Bullet 2')
doc.add_paragraph('Kỹ năng nổi bật (skills)', style='List Bullet 2')
doc.add_paragraph('Môi trường làm việc (work_environment)', style='List Bullet 2')
doc.add_paragraph('Mục tiêu nghề nghiệp (career_goal)', style='List Bullet 2')
doc.add_paragraph('Text Fields (2 trường văn bản tự do):', style='List Bullet')
doc.add_paragraph('Mô tả bản thân (self_description)', style='List Bullet 2')
doc.add_paragraph('Định hướng tương lai (future_direction)', style='List Bullet 2')

doc.add_heading('3.2 Ví Dụ Thực Tế - Học Sinh A', 2)
table2 = doc.add_table(rows=7, cols=2)
table2.style = 'Light Grid Accent 1'
table2.rows[0].cells[0].text = 'Trường'
table2.rows[0].cells[1].text = 'Giá Trị'
table2.rows[1].cells[0].text = 'Sở thích chính'
table2.rows[1].cells[1].text = 'Công nghệ'
table2.rows[2].cells[0].text = 'Môn học yêu thích'
table2.rows[2].cells[1].text = 'Tin học'
table2.rows[3].cells[0].text = 'Kỹ năng nổi bật'
table2.rows[3].cells[1].text = 'Phân tích dữ liệu'
table2.rows[4].cells[0].text = 'Mục tiêu'
table2.rows[4].cells[1].text = 'Data Scientist'
table2.rows[5].cells[0].text = 'Mô tả bản thân'
table2.rows[5].cells[1].text = 'Tôi thích giải quyết vấn đề phức tạp bằng dữ liệu'
table2.rows[6].cells[0].text = 'Định hướng tương lai'
table2.rows[6].cells[1].text = 'Muốn làm việc tại công ty công nghệ lớn'

doc.add_heading('3.3 Xử Lý Hồ Sơ Học Sinh', 2)
doc.add_paragraph('Bước 1: Chuẩn Hóa Dữ Liệu', style='List Number')
doc.add_paragraph('Chuyển toàn bộ sang không dấu, lowercase', style='List Bullet')
doc.add_paragraph('Làm sạch ký tự đặc biệt', style='List Bullet')
doc.add_paragraph('Xóa khoảng trắng thừa', style='List Bullet')
doc.add_paragraph('Bước 2: Ghép Hồ Sơ', style='List Number')
doc.add_paragraph('Nối tất cả 6 trường categorical + 2 trường text', style='List Bullet')
doc.add_paragraph('Tạo ra một đoạn text duy nhất', style='List Bullet')
doc.add_paragraph('Bước 3: Kết Quả Hồ Sơ', style='List Number')
doc.add_paragraph('"cong nghe tin hoc phan tich du lieu data scientist toi thich giai quyet van de phuc tap bang du lieu muon lam viec tai cong ty cong nghe lon"', style='List Bullet')

# Section 4
doc.add_heading('4. TF-IDF Vectorization', 1)
doc.add_paragraph('TF-IDF là kỹ thuật biến đổi văn bản thành vector số, cho phép máy tính so sánh nội dung dễ dàng hơn.')

doc.add_heading('4.1 TF (Term Frequency)', 2)
doc.add_paragraph('TF là tần suất xuất hiện của một từ trong một document.')
doc.add_paragraph('Công thức: TF(t, d) = (Số lần từ t xuất hiện trong document d) / (Tổng số từ trong document d)')
doc.add_paragraph('Ý nghĩa: Từ xuất hiện nhiều → có ý nghĩa cao hơn')

doc.add_heading('4.2 IDF (Inverse Document Frequency)', 2)
doc.add_paragraph('IDF là thước đo độ "hiếm có" của một từ trong toàn bộ corpus.')
doc.add_paragraph('Công thức: IDF(t) = log(Tổng số documents / Số documents chứa từ t)')
doc.add_paragraph('Ý nghĩa: Từ xuất hiện ít → có giá trị phân biệt cao hơn')

doc.add_heading('4.3 TF-IDF Kết Hợp', 2)
doc.add_paragraph('Công thức: TF-IDF(t, d) = TF(t, d) × IDF(t)')
doc.add_paragraph('Ý nghĩa: Từ nào vừa xuất hiện nhiều VỪA hiếm có → có trọng số cao nhất')

doc.add_heading('4.4 Kết Quả Vectorization', 2)
doc.add_paragraph('Input: Hồ sơ học sinh + Hồ sơ ngành (text)')
doc.add_paragraph('Output: Vector số (mỗi chiều = một từ khóa)', style='List Bullet')
doc.add_paragraph('Giá trị = TF-IDF weight của từ đó', style='List Bullet')
doc.add_paragraph('Dạng: Sparse matrix (lưu trữ hiệu quả)', style='List Bullet')

# Section 5
doc.add_heading('5. Cosine Similarity', 1)
doc.add_paragraph('Cosine Similarity là phương pháp toán học để đo độ tương đồng giữa 2 vector.')

doc.add_heading('5.1 Định Nghĩa Toán Học', 2)
doc.add_paragraph('Cosine Similarity(A, B) = (A · B) / (||A|| × ||B||)')
doc.add_paragraph('Trong đó:')
doc.add_paragraph('A · B = tích vô hướng của 2 vector', style='List Bullet')
doc.add_paragraph('||A|| = độ dài (norm) của vector A', style='List Bullet')
doc.add_paragraph('||B|| = độ dài (norm) của vector B', style='List Bullet')

doc.add_heading('5.2 Thang Đo & Ý Nghĩa', 2)
doc.add_paragraph('Kết quả: 0 đến 1')
doc.add_paragraph('1.0 = 2 vector giống hệt nhau (cùng hướng)', style='List Bullet')
doc.add_paragraph('0.5 = 2 vector tương đối giống', style='List Bullet')
doc.add_paragraph('0.0 = 2 vector hoàn toàn khác (vuông góc)', style='List Bullet')

doc.add_heading('5.3 Tính Chất Quan Trọng', 2)
doc.add_paragraph('Không phụ thuộc vào độ dài vector → phù hợp với TF-IDF', style='List Bullet')
doc.add_paragraph('Đo góc giữa 2 vector → so sánh định hướng, không so sánh độ lớn', style='List Bullet')
doc.add_paragraph('Nhanh & hiệu quả → phù hợp thời gian thực', style='List Bullet')

doc.add_heading('5.4 Ứng Dụng Trong Hệ Thống', 2)
doc.add_paragraph('Bước 1: Vectorize hồ sơ học sinh', style='List Number')
doc.add_paragraph('Bước 2: Vectorize tất cả 50+ hồ sơ ngành', style='List Number')
doc.add_paragraph('Bước 3: Tính Cosine Similarity giữa hồ sơ học sinh vs mỗi hồ sơ ngành', style='List Number')
doc.add_paragraph('Bước 4: Kết quả = 50+ điểm tương đồng (0-1 cho mỗi ngành)', style='List Number')

# Section 6
doc.add_heading('6. Từ Dữ Liệu Đến Điểm Số', 1)
table3 = doc.add_table(rows=6, cols=3)
table3.style = 'Light Grid Accent 1'
table3.rows[0].cells[0].text = 'Bước'
table3.rows[0].cells[1].text = 'Mô Tả'
table3.rows[0].cells[2].text = 'Kết Quả'
table3.rows[1].cells[0].text = '1'
table3.rows[1].cells[1].text = 'Chuẩn bị dữ liệu'
table3.rows[1].cells[2].text = 'Hồ sơ học sinh + Hồ sơ ngành (text)'
table3.rows[2].cells[0].text = '2'
table3.rows[2].cells[1].text = 'TF-IDF Vectorization'
table3.rows[2].cells[2].text = 'Vectors (1 × N) và (50+ × N)'
table3.rows[3].cells[0].text = '3'
table3.rows[3].cells[1].text = 'Cosine Similarity'
table3.rows[3].cells[2].text = '50+ điểm (0-1)'
table3.rows[4].cells[0].text = '4'
table3.rows[4].cells[1].text = 'Normalize'
table3.rows[4].cells[2].text = 'Điểm 0-100'
table3.rows[5].cells[0].text = '5'
table3.rows[5].cells[1].text = 'Kết Quả Cuối'
table3.rows[5].cells[2].text = 'Model Score Component'

doc.add_page_break()

# Section 7 - Conclusion
doc.add_heading('7. Kết Luận', 1)
doc.add_paragraph('Content-Based Filtering là phương pháp minh bạch, hiệu quả để so sánh hồ sơ học sinh với các ngành học.')
doc.add_paragraph('Ưu Điểm:', style='Heading 3')
doc.add_paragraph('✓ Minh bạch: Có thể giải thích tại sao một ngành được đề xuất', style='List Bullet')
doc.add_paragraph('✓ Nhanh: Không cần huấn luyện lại', style='List Bullet')
doc.add_paragraph('✓ Khách quan: Dựa trên nội dung, không phụ thuộc vào dữ liệu lịch sử', style='List Bullet')
doc.add_paragraph('Nhược Điểm:', style='Heading 3')
doc.add_paragraph('✗ Không học được patterns phức tạp từ dữ liệu', style='List Bullet')
doc.add_paragraph('✗ Phụ thuộc vào chất lượng mô tả ngành', style='List Bullet')

doc.save('CONTENT_BASED_MODULE.docx')
print("✅ Word document completed: CONTENT_BASED_MODULE.docx")
