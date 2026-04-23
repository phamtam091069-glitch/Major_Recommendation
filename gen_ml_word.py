from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('MODULE 2: MÔ HÌNH HỌC MÁY DỰ ĐOÁN (MACHINE LEARNING)', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Introduction
doc.add_heading('1. Giới Thiệu', 1)
doc.add_paragraph('Module Machine Learning là một phương pháp học từ dữ liệu lịch sử để dự đoán ngành học phù hợp cho học sinh mới.')
doc.add_paragraph('Phương pháp này hoạt động bằng cách:')
doc.add_paragraph('Thu thập dữ liệu từ hàng trăm học sinh đã chọn ngành', style='List Bullet')
doc.add_paragraph('Huấn luyện (training) một mô hình để học patterns từ dữ liệu', style='List Bullet')
doc.add_paragraph('Sử dụng mô hình đã train để dự đoán cho học sinh mới', style='List Bullet')

# Section 2
doc.add_heading('2. Dữ Liệu Huấn Luyện (Training Data)', 1)
doc.add_paragraph('Dữ liệu huấn luyện là tập hợp các học sinh đã có thông tin đầy đủ và ngành được chọn.')

doc.add_heading('2.1 Cấu Trúc Dữ Liệu', 2)
doc.add_paragraph('Mỗi dòng dữ liệu chứa:')
doc.add_paragraph('Input features (8 trường): 6 categorical + 2 text', style='List Bullet')
doc.add_paragraph('Target label: Tên ngành học (50+ lớp)', style='List Bullet')
doc.add_paragraph('Ví dụ: Nếu học sinh có sở thích "Công nghệ", kỹ năng "Lập trình" → ngành "Công nghệ thông tin"', style='List Bullet')

doc.add_heading('2.2 Số Lượng Mẫu', 2)
doc.add_paragraph('Hệ thống sử dụng ~1200 mẫu học sinh được tạo từ `data/generate_balanced_students.py`')
doc.add_paragraph('Dữ liệu cân bằng: Mỗi ngành có khoảng 80 mẫu')
doc.add_paragraph('Mục đích: Tránh mô hình bị thiên vị (bias) về ngành nào đó')

# Section 3
doc.add_heading('3. One-Hot Encoding', 1)
doc.add_paragraph('One-Hot Encoding là kỹ thuật chuyển đổi dữ liệu categorical (lựa chọn) thành dạng số để máy học có thể hiểu.')

doc.add_heading('3.1 Khái Niệm', 2)
doc.add_paragraph('Mỗi giá trị categorical được biểu diễn dưới dạng một vector 0/1')
doc.add_paragraph('Ví dụ: Sở thích chính')
doc.add_paragraph('Nếu chọn "Công nghệ" → [1, 0, 0, 0, 0, ...]', style='List Bullet')
doc.add_paragraph('Nếu chọn "Nhân văn" → [0, 1, 0, 0, 0, ...]', style='List Bullet')
doc.add_paragraph('Nếu chọn "Khoa học" → [0, 0, 1, 0, 0, ...]', style='List Bullet')

doc.add_heading('3.2 Quy Trình', 2)
doc.add_paragraph('Bước 1: Xác định tất cả giá trị có thể', style='List Number')
doc.add_paragraph('Lấy từ tất cả 1200 mẫu huấn luyện', style='List Bullet')
doc.add_paragraph('Bước 2: Tạo OneHotEncoder', style='List Number')
doc.add_paragraph('Dùng sklearn.preprocessing.OneHotEncoder', style='List Bullet')
doc.add_paragraph('Bước 3: Transform dữ liệu', style='List Number')
doc.add_paragraph('Chuyển mỗi categorical field thành vector 0/1', style='List Bullet')

doc.add_heading('3.3 Output', 2)
doc.add_paragraph('Mỗi 6 trường categorical được transform thành vector')
doc.add_paragraph('Kích thước tổng cộng: ~60-80 chiều (tùy số lượng giá trị unique)')
doc.add_paragraph('Dạng: Sparse matrix (lưu trữ hiệu quả)')

# Section 4
doc.add_heading('4. RandomForest Classifier', 1)
doc.add_paragraph('RandomForest là một thuật toán học máy mạnh mẽ, hoạt động bằng cách xây dựng nhiều cây quyết định (decision trees) và kết hợp kết quả của chúng.')

doc.add_heading('4.1 Khái Niệm Cơ Bản', 2)
doc.add_paragraph('Decision Tree: Một cây để đưa ra quyết định dựa trên các điều kiện')
doc.add_paragraph('Random Forest: Kết hợp ~100 decision trees để tăng độ chính xác')
doc.add_paragraph('Ưu điểm: Chống overfitting tốt, xử lý dữ liệu categorical tốt')

doc.add_heading('4.2 Cách Hoạt Động', 2)
doc.add_paragraph('Bước 1: Tạo nhiều subsets ngẫu nhiên từ dữ liệu training', style='List Number')
doc.add_paragraph('Bước 2: Xây dựng một decision tree cho mỗi subset', style='List Number')
doc.add_paragraph('Bước 3: Kết hợp dự đoán từ tất cả trees', style='List Number')
doc.add_paragraph('Voting: Ngành nào được dự đoán nhiều nhất sẽ được chọn', style='List Bullet')
doc.add_paragraph('Bước 4: Output probability cho mỗi ngành', style='List Number')

doc.add_heading('4.3 Trong Hệ Thống', 2)
doc.add_paragraph('File: `models/rf_model.pkl` - Mô hình RandomForest đã train')
doc.add_paragraph('Số lượng trees: ~100 (có thể điều chỉnh)')
doc.add_paragraph('Số lượng output classes: 50+ (tương ứng 50+ ngành)')

# Section 5
doc.add_heading('5. Quá Trình Huấn Luyện (Training)', 1)
doc.add_paragraph('Training là quá trình máy học từ dữ liệu để cải thiện độ chính xác.')

doc.add_heading('5.1 Dữ Liệu Huấn Luyện vs Test', 2)
doc.add_paragraph('Training set: 80% dữ liệu (~960 mẫu)')
doc.add_paragraph('Dùng để train mô hình', style='List Bullet')
doc.add_paragraph('Test set: 20% dữ liệu (~240 mẫu)')
doc.add_paragraph('Dùng để kiểm tra độ chính xác', style='List Bullet')

doc.add_heading('5.2 Quy Trình Training', 2)
doc.add_paragraph('Bước 1: Load dữ liệu từ data/raw/students.csv', style='List Number')
doc.add_paragraph('Bước 2: One-Hot Encoding 6 trường categorical', style='List Number')
doc.add_paragraph('Bước 3: TF-IDF 2 trường text', style='List Number')
doc.add_paragraph('Bước 4: Ghép thành feature matrix', style='List Number')
doc.add_paragraph('Bước 5: Train RandomForest trên feature matrix', style='List Number')
doc.add_paragraph('Bước 6: Evaluate trên test set', style='List Number')
doc.add_paragraph('Bước 7: Lưu mô hình vào models/rf_model.pkl', style='List Number')

doc.add_heading('5.3 File Script', 2)
doc.add_paragraph('Script: `train_model.py`')
doc.add_paragraph('Chạy: `python train_model.py`')
doc.add_paragraph('Output: rf_model.pkl, ohe.pkl, tfidf.pkl, evaluation.txt')

# Section 6
doc.add_heading('6. Calibration (Hiệu Chỉnh Xác Suất)', 1)
doc.add_paragraph('Calibration là kỹ thuật để làm cho xác suất từ mô hình trở nên chính xác hơn.')

doc.add_heading('6.1 Vấn Đề Mà Calibration Giải Quyết', 2)
doc.add_paragraph('Mô hình RandomForest có thể đưa ra xác suất không chính xác')
doc.add_paragraph('Ví dụ: Nếu mô hình nói xác suất 90% nhưng thực tế chỉ đúng 70%')
doc.add_paragraph('Calibration sẽ điều chỉnh để tỷ lệ xác suất phù hợp với độ đúng thực tế')

doc.add_heading('6.2 Kỹ Thuật Calibration', 2)
doc.add_paragraph('CalibratedClassifierCV: Dùng cross-validation để hiệu chỉnh')
doc.add_paragraph('Platt Scaling: Một phương pháp hiệu chỉnh đơn giản và hiệu quả')

doc.add_heading('6.3 Output Sau Calibration', 2)
doc.add_paragraph('Probability scores cho mỗi ngành: 0-1')
doc.add_paragraph('Tổng tất cả probability = 1.0')
doc.add_paragraph('Ví dụ: [0.35, 0.28, 0.15, 0.12, ..., 0.10]')

# Section 7
doc.add_heading('7. Model Score & Probability Output', 1)
doc.add_paragraph('Model Score là điểm mà mô hình RandomForest + Calibration đưa ra cho mỗi ngành.')

doc.add_heading('7.1 Output của Mô Hình', 2)
doc.add_paragraph('Vector probability: [p1, p2, p3, ..., p50+]')
doc.add_paragraph('Mỗi pi là xác suất mô hình cho ngành i')
doc.add_paragraph('Tổng: p1 + p2 + ... + p50+ = 1.0')
doc.add_paragraph('Ví dụ thực tế:', style='Heading 3')
doc.add_paragraph('Ngành "Công nghệ thông tin": 0.35 (35%)', style='List Bullet')
doc.add_paragraph('Ngành "Khoa học dữ liệu": 0.28 (28%)', style='List Bullet')
doc.add_paragraph('Ngành "Kỹ thuật phần mềm": 0.15 (15%)', style='List Bullet')
doc.add_paragraph('Các ngành khác: 0.22 (22%)', style='List Bullet')

doc.add_heading('7.2 Cách Tính Model Score', 2)
doc.add_paragraph('Model Score = Probability × 100')
doc.add_paragraph('Ví dụ: 0.35 × 100 = 35 điểm')
doc.add_paragraph('Thang điểm: 0-100')

doc.add_heading('7.3 Vai Trò Trong Hybrid Model', 2)
doc.add_paragraph('Model Score chiếm 30% trong Final Score')
doc.add_paragraph('Final Score = 0.60 × Model Score + 0.40 × Cosine Similarity')
doc.add_paragraph('Kết hợp với Content-Based để tạo kết quả tổng hợp')

# Section 8
doc.add_heading('8. Tiền Xử Lý Dữ Liệu Đầu Vào', 1)
doc.add_paragraph('Khi một học sinh mới cung cấp thông tin, dữ liệu phải được tiền xử lý giống hệt như trong quá trình training.')

doc.add_heading('8.1 Các Bước Tiền Xử Lý', 2)
doc.add_paragraph('Bước 1: Chuẩn Hóa Text', style='List Number')
doc.add_paragraph('Chuyển về không dấu, lowercase', style='List Bullet')
doc.add_paragraph('Bước 2: One-Hot Encoding', style='List Number')
doc.add_paragraph('Transform 6 trường categorical dùng OneHotEncoder đã train', style='List Bullet')
doc.add_paragraph('Bước 3: TF-IDF', style='List Number')
doc.add_paragraph('Transform 2 trường text dùng TfidfVectorizer đã train', style='List Bullet')
doc.add_paragraph('Bước 4: Ghép Features', style='List Number')
doc.add_paragraph('Nối One-Hot vectors + TF-IDF vectors', style='List Bullet')
doc.add_paragraph('Bước 5: Dự Đoán', style='List Number')
doc.add_paragraph('Đưa feature vector vào mô hình đã train để lấy probability', style='List Bullet')

doc.add_heading('8.2 Tầm Quan Trọng', 2)
doc.add_paragraph('Phải sử dụng cùng OneHotEncoder & TfidfVectorizer như training')
doc.add_paragraph('Nếu không thì kích thước vector sẽ sai')
doc.add_paragraph('Mô hình sẽ không thể dự đoán chính xác')

# Section 9
doc.add_heading('9. Đánh Giá Mô Hình (Model Evaluation)', 1)
doc.add_paragraph('Evaluation là quá trình kiểm tra xem mô hình hoạt động tốt như thế nào.')

doc.add_heading('9.1 Metrics Chính', 2)
doc.add_paragraph('Accuracy: % dự đoán đúng trên test set')
doc.add_paragraph('Precision: Trong những ngành được dự đoán là X, bao nhiêu % là đúng')
doc.add_paragraph('Recall: Trong tất cả ngành X thực tế, bao nhiêu % được dự đoán đúng')
doc.add_paragraph('F1-Score: Trung bình của Precision & Recall')

doc.add_heading('9.2 Confusion Matrix', 2)
doc.add_paragraph('Ma trận hiển thị sai lệch của mô hình cho từng ngành')
doc.add_paragraph('Dòng: Ngành thực tế')
doc.add_paragraph('Cột: Ngành được dự đoán')
doc.add_paragraph('File: `reports/confusion_matrix.csv`')

doc.add_heading('9.3 Đầu Ra Evaluation', 2)
doc.add_paragraph('File: `reports/evaluation.txt`')
doc.add_paragraph('Chứa tất cả metrics cho mỗi ngành')
doc.add_paragraph('File: `reports/per_class_metrics.csv`')
doc.add_paragraph('Metrics chi tiết cho từng ngành')

doc.add_page_break()

# Section 10 - Conclusion
doc.add_heading('10. Kết Luận', 1)
doc.add_paragraph('Module Machine Learning là thành phần học từ dữ liệu lịch sử, cung cấp thông tin xác suất dựa trên patterns.')

doc.add_heading('Ưu Điểm:', 2)
doc.add_paragraph('✓ Học được patterns phức tạp từ dữ liệu', style='List Bullet')
doc.add_paragraph('✓ RandomForest chống overfitting tốt', style='List Bullet')
doc.add_paragraph('✓ Calibration làm cho xác suất chính xác hơn', style='List Bullet')
doc.add_paragraph('✓ One-Hot Encoding xử lý categorical data tốt', style='List Bullet')

doc.add_heading('Nhược Điểm:', 2)
doc.add_paragraph('✗ Phụ thuộc vào chất lượng dữ liệu training', style='List Bullet')
doc.add_paragraph('✗ Cần retraining định kỳ khi dữ liệu thay đổi', style='List Bullet')
doc.add_paragraph('✗ Có thể gặp cold-start problem cho ngành mới', style='List Bullet')

doc.add_heading('So Sánh Với Content-Based:', 2)
table4 = doc.add_table(rows=4, cols=3)
table4.style = 'Light Grid Accent 1'
table4.rows[0].cells[0].text = 'Tiêu Chí'
table4.rows[0].cells[1].text = 'Content-Based'
table4.rows[0].cells[2].text = 'Machine Learning'
table4.rows[1].cells[0].text = 'Minh Bạch'
table4.rows[1].cells[1].text = 'Cao'
table4.rows[1].cells[2].text = 'Thấp'
table4.rows[2].cells[0].text = 'Học Pattern'
table4.rows[2].cells[1].text = 'Không'
table4.rows[2].cells[2].text = 'Có'
table4.rows[3].cells[0].text = 'Cần Training'
table4.rows[3].cells[1].text = 'Không'
table4.rows[3].cells[2].text = 'Có'

doc.save('MACHINE_LEARNING_MODULE.docx')
print("✅ Machine Learning Module completed: MACHINE_LEARNING_MODULE.docx")
