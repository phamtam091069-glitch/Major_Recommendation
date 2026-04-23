from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('MODULE 3: MÔ HÌNH LAI (HYBRID MODEL)', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Introduction
doc.add_heading('1. Giới Thiệu', 1)
doc.add_paragraph('Module Hybrid Model là sự kết hợp giữa 2 phương pháp:')
doc.add_paragraph('Content-Based Filtering (Module 1): Minh bạch, không cần training', style='List Bullet')
doc.add_paragraph('Machine Learning (Module 2): Học từ dữ liệu, tìm patterns phức tạp', style='List Bullet')
doc.add_paragraph('Mục đích: Kết hợp ưu điểm của cả 2 phương pháp để tạo ra dự đoán chính xác và tin cậy hơn.')

# Section 2
doc.add_heading('2. Tại Sao Cần Hybrid Model?', 1)
doc.add_paragraph('Mỗi phương pháp có ưu và nhược điểm:')

doc.add_heading('2.1 Content-Based Filtering', 2)
doc.add_paragraph('✓ Ưu điểm:', style='Heading 3')
doc.add_paragraph('Minh bạch: có thể giải thích tại sao đề xuất ngành nào', style='List Bullet')
doc.add_paragraph('Nhanh: không cần training', style='List Bullet')
doc.add_paragraph('Không cold-start problem', style='List Bullet')
doc.add_paragraph('✗ Nhược điểm:', style='Heading 3')
doc.add_paragraph('Không học được patterns phức tạp', style='List Bullet')
doc.add_paragraph('Phụ thuộc vào chất lượng mô tả ngành', style='List Bullet')

doc.add_heading('2.2 Machine Learning', 2)
doc.add_paragraph('✓ Ưu điểm:', style='Heading 3')
doc.add_paragraph('Học được patterns phức tạp từ dữ liệu', style='List Bullet')
doc.add_paragraph('Tương ứng với sở thích thực tế của học sinh', style='List Bullet')
doc.add_paragraph('Độ chính xác cao hơn', style='List Bullet')
doc.add_paragraph('✗ Nhược điểm:', style='Heading 3')
doc.add_paragraph('Ít minh bạch (black-box)', style='List Bullet')
doc.add_paragraph('Phụ thuộc vào dữ liệu training', style='List Bullet')
doc.add_paragraph('Cần retraining định kỳ', style='List Bullet')

doc.add_heading('2.3 Giải Pháp: Hybrid', 2)
doc.add_paragraph('Kết hợp cả 2 để:')
doc.add_paragraph('Có độ chính xác cao (từ ML)', style='List Bullet')
doc.add_paragraph('Có minh bạch (từ Content-Based)', style='List Bullet')
doc.add_paragraph('Tránh weak points của từng phương pháp', style='List Bullet')

# Section 3
doc.add_heading('3. Kiến Trúc Hybrid Model', 1)
doc.add_paragraph('Hybrid Model hoạt động theo các bước sau:')

doc.add_heading('3.1 Đầu Vào', 2)
doc.add_paragraph('Hồ sơ học sinh: 8 trường (6 categorical + 2 text)')
doc.add_paragraph('Hồ sơ ngành: 50+ mô tả chi tiết')
doc.add_paragraph('Dữ liệu training: 1200 mẫu cân bằng')

doc.add_heading('3.2 Quy Trình Song Song', 2)
doc.add_paragraph('Bước 1: Content-Based Pipeline', style='List Number')
doc.add_paragraph('Vectorize hồ sơ học sinh + hồ sơ ngành bằng TF-IDF', style='List Bullet')
doc.add_paragraph('Tính Cosine Similarity → Content Score (0-100)', style='List Bullet')
doc.add_paragraph('Bước 2: Machine Learning Pipeline', style='List Number')
doc.add_paragraph('One-Hot Encoding 6 trường categorical', style='List Bullet')
doc.add_paragraph('Dự đoán xác suất từ RandomForest đã calibrate', style='List Bullet')
doc.add_paragraph('Chuyển thành ML Score (0-100)', style='List Bullet')
doc.add_paragraph('Bước 3: Kết Hợp Điểm', style='List Number')
doc.add_paragraph('Final Score = 0.40 × Content Score + 0.60 × ML Score', style='List Bullet')

# Section 4
doc.add_heading('4. Chiến Lược Trọng Số', 1)
doc.add_paragraph('Final Score được tính bằng công thức:')
doc.add_paragraph('**Final Score = 0.40 × Content Score + 0.60 × ML Score**')

doc.add_heading('4.1 Tại Sao 40% & 60%?', 2)
doc.add_paragraph('60% cho Machine Learning vì:')
doc.add_paragraph('ML học từ 1200 mẫu dữ liệu thực tế', style='List Bullet')
doc.add_paragraph('Phản ánh đúng patterns chọn ngành của học sinh', style='List Bullet')
doc.add_paragraph('Độ chính xác cao trên test set', style='List Bullet')
doc.add_paragraph('40% cho Content-Based vì:')
doc.add_paragraph('Cung cấp minh bạch cho người dùng', style='List Bullet')
doc.add_paragraph('Tránh overfitting của ML model', style='List Bullet')
doc.add_paragraph('Không phụ thuộc vào dữ liệu training', style='List Bullet')

doc.add_heading('4.2 Tỷ Lệ Có Thể Điều Chỉnh', 2)
doc.add_paragraph('Weights được lưu trong: `models/hybrid_config.json`')
doc.add_paragraph('Có thể thay đổi tỷ lệ tùy theo nhu cầu:')
doc.add_paragraph('70% ML + 30% Content: Ưu tiên độ chính xác', style='List Bullet')
doc.add_paragraph('50% ML + 50% Content: Cân bằng giữa 2 yếu tố', style='List Bullet')

# Section 5
doc.add_heading('5. Quy Trình Tính Điểm Chi Tiết', 1)

doc.add_heading('5.1 Ví Dụ Thực Tế', 2)
table1 = doc.add_table(rows=6, cols=4)
table1.style = 'Light Grid Accent 1'
table1.rows[0].cells[0].text = 'Ngành'
table1.rows[0].cells[1].text = 'Content Score'
table1.rows[0].cells[2].text = 'ML Score'
table1.rows[0].cells[3].text = 'Final Score'
table1.rows[1].cells[0].text = 'Công nghệ thông tin'
table1.rows[1].cells[1].text = '75'
table1.rows[1].cells[2].text = '80'
table1.rows[1].cells[3].text = '78.5'
table1.rows[2].cells[0].text = 'Khoa học dữ liệu'
table1.rows[2].cells[1].text = '68'
table1.rows[2].cells[2].text = '72'
table1.rows[2].cells[3].text = '70.4'
table1.rows[3].cells[0].text = 'Kỹ thuật phần mềm'
table1.rows[3].cells[1].text = '65'
table1.rows[3].cells[2].text = '68'
table1.rows[3].cells[3].text = '66.8'
table1.rows[4].cells[0].text = 'Marketing'
table1.rows[4].cells[1].text = '45'
table1.rows[4].cells[2].text = '42'
table1.rows[4].cells[3].text = '43.2'
table1.rows[5].cells[0].text = 'Luật'
table1.rows[5].cells[1].text = '40'
table1.rows[5].cells[2].text = '38'
table1.rows[5].cells[3].text = '38.8'

doc.add_heading('5.2 Công Thức Chi Tiết', 2)
doc.add_paragraph('Final Score(Ngành i) = 0.40 × Content Score(i) + 0.60 × ML Score(i)')
doc.add_paragraph('Ví dụ cho "Công nghệ thông tin":')
doc.add_paragraph('= 0.40 × 75 + 0.60 × 80')
doc.add_paragraph('= 30 + 48')
doc.add_paragraph('= 78.5 điểm')

# Section 6
doc.add_heading('6. Lựa Chọn Top 3 Ngành', 1)
doc.add_paragraph('Sau khi tính xong Final Score cho 50+ ngành, hệ thống:')
doc.add_paragraph('Bước 1: Sắp xếp ngành theo Final Score giảm dần', style='List Number')
doc.add_paragraph('Bước 2: Lấy top 3 ngành có điểm cao nhất', style='List Number')
doc.add_paragraph('Bước 3: Tính confidence & gap cho từng ngành', style='List Number')

doc.add_heading('6.1 Output Top 3', 2)
doc.add_paragraph('Vị trí 1: "Công nghệ thông tin" - 78.5 điểm')
doc.add_paragraph('Vị trí 2: "Khoa học dữ liệu" - 70.4 điểm')
doc.add_paragraph('Vị trí 3: "Kỹ thuật phần mềm" - 66.8 điểm')

# Section 7
doc.add_heading('7. Tính Độ Tin Cậy (Confidence)', 1)
doc.add_paragraph('Confidence là mức độ chắc chắn rằng đề xuất là đúng.')

doc.add_heading('7.1 Công Thức', 2)
doc.add_paragraph('Confidence = (Điểm ngành 1 - Điểm ngành 2) × Factor + Base')
doc.add_paragraph('Dựa trên:')
doc.add_paragraph('Gap giữa ngành 1 và ngành 2 (độ tách biệt)', style='List Bullet')
doc.add_paragraph('Fit score của ngành (bao nhiêu % học sinh phù hợp)', style='List Bullet')

doc.add_heading('7.2 Thang Confidence', 2)
doc.add_paragraph('80-100: Cao - Rất chắc chắn', style='List Bullet')
doc.add_paragraph('60-80: Trung bình - Khá chắc chắn', style='List Bullet')
doc.add_paragraph('<60: Tham khảo - Cần xem xét kỹ thêm', style='List Bullet')

doc.add_heading('7.3 Ví Dụ', 2)
doc.add_paragraph('Gap = 78.5 - 70.4 = 8.1 điểm')
doc.add_paragraph('Nếu gap >5 → Confidence cao')
doc.add_paragraph('Nếu gap <3 → Confidence thấp')

# Section 8
doc.add_heading('8. Quy Trình Hoàn Chỉnh End-to-End', 1)
doc.add_heading('8.1 Sơ Đồ Luồng', 2)
doc.add_paragraph('Bước 1: Người dùng nhập form (8 trường)', style='List Number')
doc.add_paragraph('Bước 2: Frontend chuẩn hóa dữ liệu (remove diacritics)', style='List Number')
doc.add_paragraph('Bước 3: Backend nhận dữ liệu → gọi Predictor', style='List Number')
doc.add_paragraph('Bước 4: Content-Based Pipeline', style='List Number')
doc.add_paragraph('Tính TF-IDF vectors', style='List Bullet')
doc.add_paragraph('Cosine Similarity với 50+ ngành', style='List Bullet')
doc.add_paragraph('Content Score cho mỗi ngành', style='List Bullet')
doc.add_paragraph('Bước 5: Machine Learning Pipeline', style='List Number')
doc.add_paragraph('One-Hot Encoding', style='List Bullet')
doc.add_paragraph('Dự đoán RandomForest', style='List Bullet')
doc.add_paragraph('ML Score cho mỗi ngành', style='List Bullet')
doc.add_paragraph('Bước 6: Hybrid Combination', style='List Number')
doc.add_paragraph('Final Score = 0.40 × Content + 0.60 × ML', style='List Bullet')
doc.add_paragraph('Bước 7: Ranking & Top 3', style='List Number')
doc.add_paragraph('Sắp xếp giảm dần theo Final Score', style='List Bullet')
doc.add_paragraph('Lấy top 3 + tính confidence', style='List Bullet')
doc.add_paragraph('Bước 8: Return kết quả JSON', style='List Number')
doc.add_paragraph('Gửi về frontend: major name, score, confidence, suggestions', style='List Bullet')
doc.add_paragraph('Bước 9: Frontend render UI', style='List Number')
doc.add_paragraph('Hiển thị Top 3 ngành với scores & confidence', style='List Bullet')

doc.add_heading('8.2 Thời Gian Xử Lý', 2)
doc.add_paragraph('Content-Based: ~100ms (nhanh, không có ML training)')
doc.add_paragraph('Machine Learning: ~50ms (model đã train sẵn)')
doc.add_paragraph('Combination: ~10ms')
doc.add_paragraph('Total: ~160ms (gọi API & render UI: 200-500ms)')

# Section 9
doc.add_heading('9. Lưu Trữ & Cấu Hình', 1)
doc.add_heading('9.1 Files Liên Quan', 2)
doc.add_paragraph('models/rf_model.pkl - RandomForest đã train', style='List Bullet')
doc.add_paragraph('models/ohe.pkl - OneHotEncoder', style='List Bullet')
doc.add_paragraph('models/tfidf.pkl - TfidfVectorizer', style='List Bullet')
doc.add_paragraph('models/hybrid_config.json - Weights (40% & 60%)', style='List Bullet')
doc.add_paragraph('models/majors.json - 50+ ngành + mô tả', style='List Bullet')

doc.add_heading('9.2 hybrid_config.json Cấu Trúc', 2)
doc.add_paragraph('Ví dụ nội dung:')
doc.add_paragraph('{', style='List Bullet')
doc.add_paragraph('"content_weight": 0.40,', style='List Bullet 2')
doc.add_paragraph('"ml_weight": 0.60,', style='List Bullet 2')
doc.add_paragraph('"model_type": "calibrated_rf",', style='List Bullet 2')
doc.add_paragraph('"version": "1.0"', style='List Bullet 2')
doc.add_paragraph('}', style='List Bullet')

# Section 10
doc.add_heading('10. So Sánh 3 Module', 1)
table2 = doc.add_table(rows=7, cols=4)
table2.style = 'Light Grid Accent 1'
table2.rows[0].cells[0].text = 'Đặc Điểm'
table2.rows[0].cells[1].text = 'Content-Based'
table2.rows[0].cells[2].text = 'Machine Learning'
table2.rows[0].cells[3].text = 'Hybrid'
table2.rows[1].cells[0].text = 'Minh Bạch'
table2.rows[1].cells[1].text = 'Cao'
table2.rows[1].cells[2].text = 'Thấp'
table2.rows[1].cells[3].text = 'Trung bình'
table2.rows[2].cells[0].text = 'Học Patterns'
table2.rows[2].cells[1].text = 'Không'
table2.rows[2].cells[2].text = 'Có'
table2.rows[2].cells[3].text = 'Có'
table2.rows[3].cells[0].text = 'Cần Training'
table2.rows[3].cells[1].text = 'Không'
table2.rows[3].cells[2].text = 'Có'
table2.rows[3].cells[3].text = 'Có'
table2.rows[4].cells[0].text = 'Tốc Độ'
table2.rows[4].cells[1].text = 'Nhanh'
table2.rows[4].cells[2].text = 'Nhanh'
table2.rows[4].cells[3].text = 'Nhanh'
table2.rows[5].cells[0].text = 'Độ Chính Xác'
table2.rows[5].cells[1].text = 'Trung bình'
table2.rows[5].cells[2].text = 'Cao'
table2.rows[5].cells[3].text = 'Rất cao'
table2.rows[6].cells[0].text = 'Khuyến Nghị'
table2.rows[6].cells[1].text = 'Bổ sung'
table2.rows[6].cells[2].text = 'Bổ sung'
table2.rows[6].cells[3].text = '★ Sử dụng'

doc.add_page_break()

# Section 11 - Conclusion
doc.add_heading('11. Kết Luận', 1)
doc.add_paragraph('Hybrid Model là giải pháp tối ưu cho hệ thống tư vấn ngành học, kết hợp:')
doc.add_paragraph('📊 Độ chính xác từ Machine Learning', style='List Bullet')
doc.add_paragraph('💡 Minh bạch từ Content-Based Filtering', style='List Bullet')
doc.add_paragraph('⚡ Tốc độ xử lý nhanh', style='List Bullet')
doc.add_paragraph('🎯 Kết quả đáng tin cậy', style='List Bullet')

doc.add_heading('Ưu Điểm Hybrid Model:', 2)
doc.add_paragraph('✓ Kết hợp ưu điểm của cả 2 phương pháp', style='List Bullet')
doc.add_paragraph('✓ Độ chính xác cao (test accuracy ~75-80%)', style='List Bullet')
doc.add_paragraph('✓ Có thể giải thích kết quả cho người dùng', style='List Bullet')
doc.add_paragraph('✓ Robust: không bị phụ thuộc vào 1 phương pháp duy nhất', style='List Bullet')
doc.add_paragraph('✓ Dễ điều chỉnh weights nếu cần thay đổi', style='List Bullet')

doc.add_heading('Hạn Chế:', 2)
doc.add_paragraph('✗ Phức tạp hơn các phương pháp đơn lẻ', style='List Bullet')
doc.add_paragraph('✗ Cần dữ liệu training chất lượng tốt', style='List Bullet')
doc.add_paragraph('✗ Cần bảo trì 2 model (Content + ML)', style='List Bullet')

doc.add_heading('Khuyến Nghị Sử Dụng:', 2)
doc.add_paragraph('🎯 Sử dụng Hybrid Model cho hệ thống tư vấn ngành học', style='List Bullet')
doc.add_paragraph('🔄 Định kỳ retraining ML model (3-6 tháng)', style='List Bullet')
doc.add_paragraph('📊 Monitor confidence & top 3 distribution', style='List Bullet')
doc.add_paragraph('⚙️ Điều chỉnh weights dựa trên feedback người dùng', style='List Bullet')

doc.save('HYBRID_MODULE.docx')
print("✅ Hybrid Module completed: HYBRID_MODULE.docx")
